import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

INPUT_PATH = r"C:\Users\natan\OneDrive\Nati's documents\project 103 (2)\Gym_Management_System_With_TOC.docx"
OUTPUT_PATH = r"C:\Users\natan\OneDrive\Nati's documents\project 103 (2)\Gym_Management_System_Final_With_All.docx"

doc = Document(INPUT_PATH)

def insert_paragraph_after(para, text, style_name='Normal', font_name='Times New Roman', font_size=Pt(12), bold=False, alignment=None):
    """Insert a new paragraph after the given paragraph and return it."""
    new_para = doc.add_paragraph()
    
    if style_name in [s.name for s in doc.styles]:
        new_para.style = doc.styles[style_name]
    
    run = new_para.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    new_para.paragraph_format.space_before = Pt(6)
    new_para.paragraph_format.space_after = Pt(6)
    new_para.paragraph_format.line_spacing = 1.5
    
    if alignment:
        new_para.paragraph_format.alignment = alignment
    
    para._element.addnext(new_para._element)
    
    return new_para

def find_paragraph_by_text(search_text, start_idx=0):
    """Find a paragraph containing the search text."""
    for i in range(start_idx, len(doc.paragraphs)):
        if search_text.upper() in doc.paragraphs[i].text.upper():
            return doc.paragraphs[i]
    return None

# Find the LIST OF TABLES section and add LIST OF ABBREVIATIONS after it
print("Adding LIST OF ABBREVIATIONS section...")
lot_heading = find_paragraph_by_text("LIST OF TABLES")

if lot_heading:
    # Find the last table entry (Table A.1)
    last_table = find_paragraph_by_text("Table A.1 Database Schema Summary")
    
    if last_table:
        # Add empty paragraph
        current = insert_paragraph_after(last_table, '', 'Normal')
        
        # Add LIST OF ABBREVIATIONS heading
        current = insert_paragraph_after(current, 'LIST OF ABBREVIATIONS', 'Heading 1', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Add empty paragraph
        current = insert_paragraph_after(current, '', 'Normal')
        
        # Add abbreviations
        abbreviations = [
            ('ACT', 'American College of Technology'),
            ('AI', 'Artificial Intelligence'),
            ('API', 'Application Programming Interface'),
            ('BMI', 'Body Mass Index'),
            ('CI/CD', 'Continuous Integration / Continuous Deployment'),
            ('CSRF', 'Cross-Site Request Forgery'),
            ('CSS', 'Cascading Style Sheets'),
            ('DB', 'Database'),
            ('ETB', 'Ethiopian Birr'),
            ('GMS', 'Gym Management System'),
            ('GPS', 'Global Positioning System'),
            ('HTML', 'HyperText Markup Language'),
            ('HTTP', 'HyperText Transfer Protocol'),
            ('HTTPS', 'HyperText Transfer Protocol Secure'),
            ('IDE', 'Integrated Development Environment'),
            ('JS', 'JavaScript'),
            ('MVT', 'Model-View-Template'),
            ('OOSAD', 'Object-Oriented System Analysis and Design'),
            ('ORM', 'Object-Relational Mapping'),
            ('OS', 'Operating System'),
            ('PDF', 'Portable Document Format'),
            ('QR', 'Quick Response'),
            ('RAM', 'Random Access Memory'),
            ('RWD', 'Responsive Web Design'),
            ('SDK', 'Software Development Kit'),
            ('SQL', 'Structured Query Language'),
            ('SSD', 'Solid State Drive'),
            ('SSL', 'Secure Sockets Layer'),
            ('UI', 'User Interface'),
            ('URL', 'Uniform Resource Locator'),
            ('VM', 'Virtual Machine'),
            ('XSS', 'Cross-Site Scripting'),
        ]
        
        for abbr, full in abbreviations:
            new_para = doc.add_paragraph()
            
            r1 = new_para.add_run(f'{abbr}')
            r1.font.name = 'Times New Roman'
            r1.font.size = Pt(12)
            r1.font.bold = True
            r1.font.color.rgb = RGBColor(0, 0, 0)
            
            r2 = new_para.add_run(f'  -  {full}')
            r2.font.name = 'Times New Roman'
            r2.font.size = Pt(12)
            r2.font.color.rgb = RGBColor(0, 0, 0)
            
            new_para.paragraph_format.space_before = Pt(2)
            new_para.paragraph_format.space_after = Pt(2)
            new_para.paragraph_format.line_spacing = 1.5
            
            current._element.addnext(new_para._element)
            current = new_para
        
        print(f"  Added LIST OF ABBREVIATIONS section with {len(abbreviations)} entries")
    else:
        print("  WARNING: Could not find last table entry")
else:
    print("  WARNING: Could not find LIST OF TABLES heading")

doc.save(OUTPUT_PATH)
print(f'\nDocument saved: {OUTPUT_PATH}')
print('LIST OF ABBREVIATIONS section added successfully.')
