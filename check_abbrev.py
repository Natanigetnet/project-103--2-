import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r"C:\Users\natan\OneDrive\Nati's documents\project 103 (2)\Gym_Management_System_With_TOC.docx")

print("=== Searching for LIST OF ABBREVIATIONS ===")
for i, para in enumerate(doc.paragraphs):
    if 'ABBREVIATION' in para.text.upper():
        print(f'{i}: [{para.style.name}] |{para.text}|')
        # Show next 5 paragraphs
        for j in range(i+1, min(i+6, len(doc.paragraphs))):
            print(f'  {j}: [{doc.paragraphs[j].style.name}] |{doc.paragraphs[j].text[:50]}|')
        break

print("\n=== TABLE OF CONTENTS (Chapters 5-6) ===")
for i, para in enumerate(doc.paragraphs):
    if 'CHAPTER FIVE' in para.text.upper():
        for j in range(i, min(i+30, len(doc.paragraphs))):
            if doc.paragraphs[j].text.strip():
                print(f'{j}: {doc.paragraphs[j].text[:100]}')
        break

print("\n=== LIST OF TABLES (new entries) ===")
for i, para in enumerate(doc.paragraphs):
    if 'Table 5.1' in para.text:
        for j in range(max(0, i-2), min(i+8, len(doc.paragraphs))):
            if doc.paragraphs[j].text.strip():
                print(f'{j}: {doc.paragraphs[j].text[:100]}')
        break
