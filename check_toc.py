import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r"C:\Users\natan\OneDrive\Nati's documents\project 103 (2)\Gym_Management_System_Updated_Ch5_6.docx")

# Find the TOC, List of Figures, List of Tables, List of Abbreviations sections
print("=== Looking for TOC, List of Figures, List of Tables, List of Abbreviations ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip().upper()
    if any(kw in text for kw in ['TABLE OF CONTENT', 'LIST OF FIGURE', 'LIST OF TABLE', 'LIST OF ABBREV', 'HEADING 1', 'HEADING 2']):
        if any(kw in text for kw in ['TABLE OF CONTENT', 'LIST OF FIGURE', 'LIST OF TABLE', 'LIST OF ABBREV']):
            print(f'{i}: [{para.style.name}] {para.text[:120]}')
        elif 'CHAPTER' in text or 'FIGURE' in text or 'TABLE' in text:
            print(f'{i}: [{para.style.name}] {para.text[:120]}')

print("\n=== Paragraphs 35-70 (TOC area) ===")
for i in range(35, 70):
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        if p.text.strip():
            print(f'{i}: [{p.style.name}] {p.text[:120]}')

print("\n=== All tables with captions ===")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip().startswith('Table'):
        print(f'{i}: [{para.style.name}] {para.text[:120]}')

print("\n=== All figure captions ===")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip().startswith('Figure'):
        print(f'{i}: [{para.style.name}] {para.text[:120]}')
