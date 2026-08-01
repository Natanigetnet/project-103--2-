import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r"C:\Users\natan\OneDrive\Nati's documents\project 103 (2)\Gym_Management_System_Updated_Ch5_6.docx")
print(f'Total paragraphs: {len(doc.paragraphs)}')
print(f'Total tables: {len(doc.tables)}')

# Check images preserved
body = doc.element.body
imgdata = body.findall('.//{urn:schemas-microsoft-com:vml}imagedata')
print(f'VML images preserved: {len(imgdata)}')

# Count chapter headings
for i, para in enumerate(doc.paragraphs):
    if 'CHAPTER' in para.text.upper() and ('FIVE' in para.text.upper() or 'SIX' in para.text.upper()):
        print(f'Para {i}: {para.text[:80]}')

# Count approximate pages
total_chars = sum(len(p.text) for p in doc.paragraphs)
print(f'Total characters: {total_chars}')
print(f'Estimated pages (chars/3000): {total_chars / 3000:.1f}')

# Show last 30 paragraphs with content
print('\n--- Last 30 content paragraphs ---')
content_paras = [(i, p) for i, p in enumerate(doc.paragraphs) if p.text.strip()]
for i, p in content_paras[-30:]:
    print(f'{i}: [{p.style.name}] {p.text[:100]}')
