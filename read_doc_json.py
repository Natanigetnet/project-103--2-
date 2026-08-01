from docx import Document
import json

doc = Document(r"C:\Users\natan\OneDrive\Nati's documents\project 103 (2)\Gym_Management_System_123.docx")

content = {
    'paragraphs': [],
    'tables': []
}

for i, p in enumerate(doc.paragraphs):
    content['paragraphs'].append({
        'index': i,
        'text': p.text,
        'style': p.style.name
    })

for i, table in enumerate(doc.tables):
    table_data = []
    for row in table.rows:
        row_data = [cell.text for cell in row.cells]
        table_data.append(row_data)
    content['tables'].append({
        'index': i,
        'rows': table_data
    })

with open(r"C:\Users\natan\OneDrive\Nati's documents\project 103 (2)\doc_content.json", 'w', encoding='utf-8') as f:
    json.dump(content, f, indent=2, ensure_ascii=False)

print("Content saved to doc_content.json")
