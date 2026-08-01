from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_table_with_style(doc, data, headers=None):
    """Helper to add a formatted table"""
    if headers:
        # Create table with header row + data rows
        table = doc.add_table(rows=1 + len(data), cols=len(headers))
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            # Make header bold
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Add data rows
        for row_idx, row_data in enumerate(data, start=1):
            row_cells = table.rows[row_idx].cells
            for col_idx, cell_data in enumerate(row_data):
                row_cells[col_idx].text = str(cell_data)
    else:
        table = doc.add_table(rows=len(data), cols=len(data[0]))
        # Add data rows
        for row_idx, row_data in enumerate(data):
            row_cells = table.rows[row_idx].cells
            for col_idx, cell_data in enumerate(row_data):
                row_cells[col_idx].text = str(cell_data)
    
    # Add borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)
    tblPr.append(borders)
    
    return table

# Create document
doc = Document()

# Title
title = doc.add_heading('Future Gym Management System', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Complete System Documentation')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(14)
subtitle_format.font.italic = True

doc.add_paragraph('Project 103 - Django Web Application')
doc.add_paragraph()

# Table of Contents
doc.add_heading('Table of Contents', 1)
toc_items = [
    '1. System Overview',
    '2. Technology Stack',
    '3. Project Structure',
    '4. User Roles & Permissions',
    '5. Database Models',
    '6. Authentication & Authorization',
    '7. Key Features',
    '8. URL Routes',
    '9. Notification System',
    '10. Running the System',
    '11. Security Features',
    '12. Database Relationships',
    '13. Sample Data',
    '14. Environment Variables'
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# Section 1: System Overview
doc.add_heading('1. System Overview', 1)
doc.add_paragraph(
    'Future Gym is a comprehensive web-based gym management application built with Django (Python). '
    'It manages all aspects of a fitness center including member registration, trainer assignments, '
    'attendance tracking via QR codes, training plans, payments, scheduling, AI chatbot, social feed, '
    'and more.'
)
doc.add_paragraph(
    'The system is designed for deployment on Render (cloud hosting) with Cloudinary for image storage '
    'and SQLite (locally) or PostgreSQL (production) as the database.'
)

# Section 2: Technology Stack
doc.add_heading('2. Technology Stack', 1)
tech_data = [
    ['Python 3.x', 'Programming language', 'Core language for all logic'],
    ['Django 5.2+', 'Web framework', 'URL routing, ORM, templates, auth, admin'],
    ['SQLite/PostgreSQL', 'Database', 'Data storage (local/prod)'],
    ['HTML/CSS/Bootstrap 5.3', 'Frontend', 'Responsive layouts, glassmorphism design'],
    ['Pillow', 'Image library', 'Profile pictures, QR codes'],
    ['qrcode', 'QR generator', 'Member ID card QR codes'],
    ['Google Gemini AI', 'AI API', 'Chat assistant (gemini-2.0-flash)'],
    ['Cloudinary', 'Cloud storage', 'Image hosting in production'],
    ['python-dotenv', 'Env variables', 'Secrets management'],
    ['gunicorn', 'WSGI server', 'Production server on Render'],
    ['Telegram Bot API', 'Messaging', 'Group broadcasts'],
    ['Chapa', 'Payment gateway', 'Ethiopian payment processor'],
    ['Django Email (SMTP)', 'Email system', 'Notifications via Gmail']
]
add_table_with_style(doc, tech_data, ['Technology', 'Type', 'Purpose'])

doc.add_page_break()

# Section 3: Project Structure
doc.add_heading('3. Project Structure', 1)
structure_text = """project 103 (2)/
├── data.json                    # Database fixture (all records as JSON)
├── SYSTEM_DOCUMENTATION.md      # This documentation
├── IMPLEMENTATION_SUMMARY.md    # Developer notes
├── QUICK_REFERENCE.md           # Quick reference card
├── venv/                        # Python virtual environment
├── .gitignore                   # Git ignore rules
└── project 103/
    ├── TRAINER_CATEGORIES_GUIDE.md
    └── blog/                    # Django project root
        ├── manage.py            # Django CLI tool
        ├── requirements.txt     # Python dependencies
        ├── db.sqlite3           # Local database
        ├── .env                 # Environment variables
        ├── media/               # Uploaded files
        ├── blog/                # Project config
        │   ├── settings.py      # All configuration
        │   ├── urls.py          # URL routing (~127 routes)
        │   ├── wsgi.py          # Production entry point
        │   └── asgi.py          # Async entry point
        └── news/                # Main app (business logic)
            ├── models.py        # 20+ database models
            ├── views.py         # ~4196 lines of logic
            ├── forms.py         # Form definitions
            ├── admin.py         # Admin panel config
            ├── email_utils.py   # Email utilities
            ├── migrations/      # 48 schema migrations
            └── templates/       # 80+ HTML templates"""
doc.add_paragraph(structure_text)

doc.add_page_break()

# Section 4: User Roles
doc.add_heading('4. User Roles & Permissions', 1)

doc.add_heading('4.1 Admin (Superuser)', 2)
doc.add_paragraph('Gym owner/manager with full system access.')
doc.add_paragraph('Key capabilities:', style='List Bullet')
admin_caps = [
    'Manage all users and data',
    'Configure gym settings',
    'Record payments and view financial reports',
    'Manage trainer schedules and employee payments',
    'Scan QR codes for attendance',
    'Broadcast Telegram messages',
    'Manage social feed'
]
for cap in admin_caps:
    doc.add_paragraph(cap, style='List Bullet 2')
doc.add_paragraph('Key pages: /admin-portal/, /admin/, /manage-users/, /members/, /record-payment/, /admin/income-report/')

doc.add_heading('4.2 Trainer', 2)
doc.add_paragraph('Fitness instructors assigned to train members.')
trainer_caps = [
    'View assigned trainees',
    'Create training sessions',
    'Create training plans with exercises',
    'View trainee BMI/progress',
    'View work schedule',
    'See which trainees are in the gym',
    'Receive ratings from trainees'
]
for cap in trainer_caps:
    doc.add_paragraph(cap, style='List Bullet')
doc.add_paragraph('Key pages: /tracker/, /session/create/, /training-plan/<id>/, /trainer/workout-tracking/')

doc.add_heading('4.3 Trainee (Member)', 2)
doc.add_paragraph('Regular gym members/customers.')
trainee_caps = [
    'View profile and track BMI',
    'Register for training sessions',
    'View training plan/schedule',
    'Rate trainer and request changes',
    'View digital ID card with QR code',
    'View social feed',
    'Make membership payments',
    'Chat with AI assistant'
]
for cap in trainee_caps:
    doc.add_paragraph(cap, style='List Bullet')
doc.add_paragraph('Key pages: /home/, /settings/, /bmi/, /trainer-sessions/, /my-id-card/, /feed/')

doc.add_heading('4.4 Registrar', 2)
doc.add_paragraph('Front desk staff for check-ins and registration.')
registrar_caps = [
    'Register new trainees',
    'Scan QR codes for check-in/check-out',
    'View who is currently in the gym',
    'View attendance logs'
]
for cap in registrar_caps:
    doc.add_paragraph(cap, style='List Bullet')
doc.add_paragraph('Key pages: /registrar/dashboard/, /registrar/register/, /registrar/scan-qr/')

doc.add_page_break()

# Section 5: Database Models
doc.add_heading('5. Database Models', 1)
doc.add_paragraph('The system uses 20+ database models to store all data. Key models include:')

doc.add_heading('5.1 Core User/Member Models', 2)

doc.add_heading('Category', 3)
doc.add_paragraph('Training types offered at the gym.')
doc.add_paragraph('Fields: name, description')
doc.add_paragraph('Example: {name: "Aerobics", description: "High-energy fitness classes..."}')

doc.add_heading('names', 3)
doc.add_paragraph('Central member/person record. Every trainer, trainee, and member has a row here.')
doc.add_paragraph('Fields: name, email, phone_number, detail, date, image, role (trainer/trainee), gender, trainer (FK to User), preferred_trainer (FK to User), category (FK to Category)')
doc.add_paragraph('Example: {name: "kyle", email: "kylen@gmail.com", role: "trainee", trainer: [john], category: Aerobics}')

doc.add_heading('UserProfile', 3)
doc.add_paragraph('Extended profile linked 1-to-1 with Django User model.')
doc.add_paragraph('Fields: user (1-to-1), role (trainer/trainee/registrar), gender, category (FK), medical_info, image, created_at')

doc.add_heading('5.2 Messaging Models', 2)

doc.add_heading('questions', 3)
doc.add_paragraph('Incoming messages/questions (also used as notification headers).')
doc.add_paragraph('Fields: name (sender), email, quest (message text), ai_answered (boolean)')

doc.add_heading('response_model', 3)
doc.add_paragraph('Replies to questions (also used as notification bodies).')
doc.add_paragraph('Fields: name (FK to User), quest (FK to questions), text, is_read (boolean)')

doc.add_heading('5.3 Payment Models', 2)

doc.add_heading('MembershipPayment', 3)
doc.add_paragraph('Records membership fee payments from trainees.')
doc.add_paragraph('Fields: user (FK), amount, payment_date, entry_date, payment_method (CASH/TRANSFER/CHAPA/OTHER), receipt_number, is_verified, subscription_start, subscription_end, chapa_tx_ref')

doc.add_heading('TrainerPayment', 3)
doc.add_paragraph('Salary/payment info for trainers and employees.')
doc.add_paragraph('Fields: trainer (1-to-1), salary, last_payment_date, payment_frequency (weekly/biweekly/monthly), notes')
doc.add_paragraph('Properties: next_payment_due, days_until_payment')

doc.add_heading('5.4 Fitness/Training Models', 2)

doc.add_heading('BodyMetric', 3)
doc.add_paragraph('BMI and body measurements tracked over time.')
doc.add_paragraph('Fields: user (FK), weight (kg), height (cm), bmi (auto-calculated), recorded_at')
doc.add_paragraph('Auto-calculation: bmi = weight / (height_in_meters ^ 2)')

doc.add_heading('TrainingSession', 3)
doc.add_paragraph('Scheduled class/session created by a trainer.')
doc.add_paragraph('Fields: title, description, session_date, duration_minutes, space (FK), trainer (FK), max_trainees, registered_trainees (M2M), created_at')
doc.add_paragraph('Properties: is_full, slots_left, end_time, is_past')

doc.add_heading('TrainingPlan', 3)
doc.add_paragraph('Structured workout program created by trainer for trainee.')
doc.add_paragraph('Fields: trainee (FK), trainer (FK), split_type, start_date, end_date, notes, is_active')
doc.add_paragraph('Split types: upper_lower, push_pull_legs, leg_arm_chest_back, full_body, bro_split, custom')

doc.add_heading('TrainingPlanDay', 3)
doc.add_paragraph('Individual days within a training plan.')
doc.add_paragraph('Fields: plan (FK), day_index, day_label, is_rest_day, exercises (JSON)')
doc.add_paragraph('Exercises format: [{name, sets, reps, weight, notes}, ...]')

doc.add_heading('5.5 Attendance/ID Models', 2)

doc.add_heading('MemberID', 3)
doc.add_paragraph('Digital ID card with QR code for each member.')
doc.add_paragraph('Fields: member (1-to-1), unique_id (UUID), qr_code (image), created_at')

doc.add_heading('AttendanceLog', 3)
doc.add_paragraph('Check-in/check-out records.')
doc.add_paragraph('Fields: member (FK), session (FK), checked_in_by (FK), check_in, check_out, checked_out_by (FK), notes')

doc.add_heading('5.6 Social/Feedback Models', 2)

doc.add_heading('TrainerRating', 3)
doc.add_paragraph('Trainee ratings of trainers (0-5 stars).')
doc.add_paragraph('Fields: trainee (FK), trainer (FK), rating (0-5), comment, created_at')
doc.add_paragraph('Constraint: unique_together (trainee, trainer)')

doc.add_heading('FeedPost', 3)
doc.add_paragraph('Social media-style posts in the gym feed.')
doc.add_paragraph('Fields: author (FK), image, quote, hashtags, hyped_by (M2M), created_at')
doc.add_paragraph('Properties: hype_count, image_url')

doc.add_page_break()

# Section 6: Authentication
doc.add_heading('6. Authentication & Authorization', 1)

doc.add_heading('6.1 Login System', 2)
doc.add_paragraph('Uses Django built-in django.contrib.auth system.')
doc.add_paragraph('Case-insensitive login: Custom backend allows username in any case (e.g., "John" matches "john").')
doc.add_paragraph('Login page at /login/. First login prompts password change for non-superusers.')
doc.add_paragraph('Sessions stored in database.')

doc.add_heading('6.2 Access Control Decorators', 2)
doc.add_paragraph('@user_passes_test(lambda u: u.is_superuser) - Admin only (~30 views)')
doc.add_paragraph('@login_required - Any authenticated user')
doc.add_paragraph('@_require_trainee - Ensures UserProfile.role == "trainee"')
doc.add_paragraph('@_require_trainer - Ensures UserProfile.role == "trainer"')
doc.add_paragraph('@_require_registrar - Ensures UserProfile.role == "registrar" (superusers also allowed)')

doc.add_heading('6.3 Signup Flow', 2)
doc.add_paragraph('1. User visits /signup/ and fills form (username, email, password, name, phone, gender, role, category)')
doc.add_paragraph('2. Admin creating account: all fields required including role and category')
doc.add_paragraph('3. Django User created + UserProfile + optionally names record')
doc.add_paragraph('4. Welcome email sent via send_email_async()')
doc.add_paragraph('5. User automatically logged in and redirected to /home/')

doc.add_page_break()

# Section 7: Key Features
doc.add_heading('7. Key Features', 1)

doc.add_heading('7.1 QR Code Attendance System', 2)
doc.add_paragraph('Purpose: Track when members arrive at and leave the gym.')
doc.add_paragraph('Workflow:')
doc.add_paragraph('1. ID Generation: Member views /my-id-card/, system creates MemberID with UUID and generates QR code', style='List Number')
doc.add_paragraph('2. Check-In: Registrar/admin scans QR at /registrar/scan-qr/ or /admin/qr-checkin/, POST to /api/check-in/', style='List Number')
doc.add_paragraph('3. Check-Out: Same process to /api/check-out/, sets check_out = now()', style='List Number')
doc.add_paragraph('4. Toggle Logic: /api/record-attendance/ checks OUT if open session exists, otherwise checks IN', style='List Number')
doc.add_paragraph('5. Split Advancement: On check-out, trainee SplitProgression advances to next day (once per calendar day)', style='List Number')
doc.add_paragraph('6. Membership Check: _has_active_membership() verifies subscription_end >= today', style='List Number')

doc.add_heading('7.2 Training Plan System', 2)
doc.add_paragraph('Purpose: Trainers create structured workout programs for trainees.')
doc.add_paragraph('Workflow:')
doc.add_paragraph('1. Create Plan: Trainer goes to /training-plan/<trainee_id>/, selects split type', style='List Number')
doc.add_paragraph('2. Add Exercises: For each day, add exercises with name, sets, reps, weight, notes (stored as JSON)', style='List Number')
doc.add_paragraph('3. Weekly Calendar: 7-day view with week navigation, shows label, exercises, rest days', style='List Number')
doc.add_paragraph('4. Progression Tracking: SplitProgression tracks current day in cycle, advances on check-out', style='List Number')
doc.add_paragraph('5. Notifications: Trainee notified via questions/response_model when plan created/modified', style='List Number')

doc.add_heading('7.3 AI Chat System', 2)
doc.add_paragraph('Purpose: Intelligent assistant for gym-related questions.')
doc.add_paragraph('Workflow:')
doc.add_paragraph('1. User visits /chat/ and types question', style='List Number')
doc.add_paragraph('2. JavaScript POSTs to /api/chat/ with message and history (last 10 messages)', style='List Number')
doc.add_paragraph('3. Server builds gym context: categories, trainers (with ratings), sessions, spaces, hours, stats', style='List Number')
doc.add_paragraph('4. Builds site guide: detailed instructions for every page/feature by role', style='List Number')
doc.add_paragraph('5. Local FAQ matcher: 25+ regex patterns for instant answers (e.g., gym hours)', style='List Number')
doc.add_paragraph('6. If no local match: calls Google Gemini API (gemini-2.0-flash) with system prompt', style='List Number')
doc.add_paragraph('7. Conversation saved to questions/response_model for admin review', style='List Number')

doc.add_heading('7.4 Payment & Subscription System', 2)
doc.add_paragraph('Purpose: Track membership payments and subscription validity.')
doc.add_paragraph('Workflow:')
doc.add_paragraph('1. Admin Records Payment: /record-payment/, creates MembershipPayment + notification', style='List Number')
doc.add_paragraph('2. Self-Service via Chapa: /membership/pay/, fee 5000 ETB for 3 months, creates verified payment', style='List Number')
doc.add_paragraph('3. Subscription Alerts: Home page checks latest payment, shows "expiring" (7 days) or "expired"', style='List Number')
doc.add_paragraph('4. Income Report: /admin/income-report/ calculates gross income, expenses, tax (30%), net income', style='List Number')

doc.add_heading('7.5 Employee Payment & Scheduling', 2)
doc.add_paragraph('Purpose: Manage trainer salaries and work schedules.')
doc.add_paragraph('Workflow:')
doc.add_paragraph('1. Gym Config: /admin/gym-config/ sets global payment day (1-28)', style='List Number')
doc.add_paragraph('2. Employee Payments: /admin/employee-payments/ sets salary, frequency, calculates next_payment_due', style='List Number')
doc.add_paragraph('3. Trainer Schedules: /admin/trainer-schedules/ assigns days/shifts (day: 06:00-16:00, evening: 16:00-23:00)', style='List Number')
doc.add_paragraph('4. Trainer Dashboard: /admin/trainer-dashboard/ shows monthly hours, attendance rate, calendar', style='List Number')
doc.add_paragraph('5. Trainer Schedule: /trainer/my-schedule/ shows 4-week calendar, add comments sent to admin', style='List Number')

doc.add_heading('7.6 Social Feed', 2)
doc.add_paragraph('Purpose: Social media-style feed for gym community.')
doc.add_paragraph('Workflow:')
doc.add_paragraph('1. Create posts at /feed/create/ with image, quote, hashtags', style='List Number')
doc.add_paragraph('2. Posts at /feed/ in reverse chronological order, paginated (6 per page)', style='List Number')
doc.add_paragraph('3. "Hype" (like) posts - toggle stored in FeedPost.hyped_by (M2M)', style='List Number')
doc.add_paragraph('4. Authors can edit/delete own posts. Admin can remove any post via /feed/admin/', style='List Number')
doc.add_paragraph('5. Hashtags parsed from space/comma-separated input, normalized with # prefix', style='List Number')

doc.add_heading('7.7 Telegram Integration', 2)
doc.add_paragraph('Purpose: Broadcast messages to gym Telegram group.')
doc.add_paragraph('Workflow:')
doc.add_paragraph('1. Admin goes to /telegram-broadcast/, types message + optional image', style='List Number')
doc.add_paragraph('2. System calls Telegram Bot API (sendPhoto or sendMessage) with bot token and chat ID', style='List Number')
doc.add_paragraph('3. Message appears in Telegram group instantly', style='List Number')
doc.add_paragraph('4. Trainer session creation can also send Telegram notification', style='List Number')

doc.add_page_break()

# Section 8: URL Routes
doc.add_heading('8. URL Routes (All ~127 Endpoints)', 1)
doc.add_paragraph('Complete routing table from blog/urls.py:')

url_data = [
    ['/', 'landing', 'Anyone', 'Public landing page'],
    ['/home/', 'home', 'Logged in', 'Main dashboard'],
    ['/login/', 'loginUser', 'Anyone', 'Login page'],
    ['/signup/', 'signup', 'Anyone', 'Registration page'],
    ['/logout/', 'logoutUser', 'Logged in', 'Log out'],
    ['/members/', 'members', 'Admin', 'View all members'],
    ['/register/', 'register', 'Admin', 'Register new member'],
    ['/detail/<name>', 'detail', 'Logged in', 'Member detail page'],
    ['/category_list/', 'category_list', 'Admin', 'Manage categories'],
    ['/record-payment/', 'record_payment', 'Admin', 'Record payment'],
    ['/tracker/', 'trainer_tracker', 'Trainer', 'Trainer trainee list'],
    ['/assign-trainer/', 'assign_trainer', 'Admin', 'Assign trainers'],
    ['/settings/', 'trainee_settings', 'Trainee', 'Trainee settings'],
    ['/session/create/', 'create_session', 'Trainer', 'Create session'],
    ['/bmi/', 'trainee_bmi', 'Trainee', 'BMI tracker'],
    ['/my-id-card/', 'my_id_card', 'Logged in', 'Digital ID card'],
    ['/chat/', 'chat_page', 'Logged in', 'AI chat page'],
    ['/training-plan/<id>/', 'training_plan_view', 'Trainer/Trainee', 'Training plan'],
    ['/feed/', 'feed_view', 'Logged in', 'Social feed'],
    ['/registrar/dashboard/', 'registrar_dashboard', 'Registrar', 'Registrar dashboard'],
    ['/rate-trainer/<name>/', 'rate_trainer', 'Trainee', 'Rate trainer'],
    ['/membership/pay/', 'membership_payment_page', 'Logged in', 'Payment page'],
    ['/admin-portal/', 'admin_dash', 'Admin', 'Admin portal'],
    ['/admin/gym-config/', 'gym_config_view', 'Admin', 'Gym settings'],
    ['/admin/trainer-dashboard/', 'admin_trainer_dashboard', 'Admin', 'Trainer analytics'],
    ['/admin/income-report/', 'income_report', 'Admin', 'Financial report'],
]
add_table_with_style(doc, url_data, ['URL', 'View', 'Access', 'Purpose'])

doc.add_page_break()

# Section 9: Notification System
doc.add_heading('9. Notification System', 1)
doc.add_paragraph('The system repurposes questions + response_model as a generic notification system:')
doc.add_paragraph('1. When something happens (trainer assignment, session creation, payment, etc.), system creates questions record (header) + response_model record (body)', style='List Number')
doc.add_paragraph('2. The name field on response_model points to User who should see notification', style='List Number')
doc.add_paragraph('3. The is_read field tracks whether user has seen it', style='List Number')
doc.add_paragraph('4. Home page counts unread notifications and displays badge', style='List Number')
doc.add_paragraph('5. Visiting /response_list/ or /questions-log/ marks unread as read', style='List Number')
doc.add_paragraph('Example: Trainer creates session -> For each trainee, questions created with quest="New training session: bb" and response_model with full details + registration link -> Trainee sees unread badge -> Clicks "Messages" to view')

# Section 10: Running the System
doc.add_heading('10. Running the System', 1)

doc.add_heading('10.1 Local Development', 2)
doc.add_paragraph('# 1. Activate virtual environment')
doc.add_paragraph('venv\\Scripts\\activate')
doc.add_paragraph('# 2. Install dependencies')
doc.add_paragraph('pip install -r requirements.txt')
doc.add_paragraph('# 3. Set up environment variables in .env file')
doc.add_paragraph('# DJANGO_SECRET_KEY, EMAIL_HOST_USER, EMAIL_HOST_PASSWORD, GEMINI_API_KEY, etc.')
doc.add_paragraph('# 4. Run migrations (create database tables)')
doc.add_paragraph('python manage.py migrate')
doc.add_paragraph('# 5. Load sample data (optional)')
doc.add_paragraph('python manage.py loaddata data.json')
doc.add_paragraph('# 6. Create admin user')
doc.add_paragraph('python manage.py createsuperuser')
doc.add_paragraph('# 7. Start development server')
doc.add_paragraph('python manage.py runserver')
doc.add_paragraph('Then visit http://localhost:8000/')

doc.add_heading('10.2 Production (Render)', 2)
doc.add_paragraph('Uses gunicorn as WSGI server')
doc.add_paragraph('Database is PostgreSQL via dj-database-url (DATABASE_URL env var)')
doc.add_paragraph('Images stored on Cloudinary')
doc.add_paragraph('Static files collected with: python manage.py collectstatic')

# Section 11: Security
doc.add_heading('11. Security Features', 1)
security_data = [
    ['CSRF Protection', 'Django middleware adds CSRF tokens to all forms. POST without valid token rejected.'],
    ['Password Hashing', 'PBKDF2-SHA256 with 1,200,000 iterations'],
    ['Session Management', 'Django sessions stored in database with expiry dates'],
    ['Clickjacking Protection', 'XFrameOptionsMiddleware prevents iframe embedding'],
    ['Role-Based Access', 'Decorators ensure users only access appropriate pages'],
    ['Secret Keys', 'Loaded from environment variables, not hardcoded'],
    ['Email Verification', 'Welcome emails sent on registration with credentials']
]
add_table_with_style(doc, security_data, ['Feature', 'How It Works'])

doc.add_page_break()

# Section 12: Database Relationships
doc.add_heading('12. Database Relationships', 1)
doc.add_paragraph('User (Django built-in)')
doc.add_paragraph(' ├── UserProfile (1-to-1) -- role, gender, medical_info, image', style='List Bullet')
doc.add_paragraph(' ├── MembershipPayment (1-to-many) -- amount, method, subscription dates', style='List Bullet')
doc.add_paragraph(' ├── BodyMetric (1-to-many) -- weight, height, bmi', style='List Bullet')
doc.add_paragraph(' ├── FeedPost (1-to-many) -- image, quote, hashtags', style='List Bullet')
doc.add_paragraph(' └── response_model (1-to-many) -- responses received', style='List Bullet')

doc.add_paragraph('names (central member record)')
doc.add_paragraph(' ├── trainer (FK to User) -- who trains this person', style='List Bullet')
doc.add_paragraph(' ├── category (FK to Category) -- training type', style='List Bullet')
doc.add_paragraph(' ├── MemberID (1-to-1) -- UUID + QR code', style='List Bullet')
doc.add_paragraph(' ├── AttendanceLog (1-to-many) -- check-in/out records', style='List Bullet')
doc.add_paragraph(' ├── TrainingSession (as trainer, 1-to-many) -- sessions they lead', style='List Bullet')
doc.add_paragraph(' ├── TrainingSession (as trainee, M2M) -- sessions they registered for', style='List Bullet')
doc.add_paragraph(' ├── TrainerRating (as trainee/trainer, 1-to-many)', style='List Bullet')
doc.add_paragraph(' ├── TrainingPlan (as trainee/trainer, 1-to-many)', style='List Bullet')
doc.add_paragraph(' ├── SplitProgression (1-to-1) -- split cycle tracking', style='List Bullet')
doc.add_paragraph(' ├── TrainerPayment (1-to-1) -- salary info', style='List Bullet')
doc.add_paragraph(' └── TrainerSchedule (1-to-many) -- weekly work schedule', style='List Bullet')

doc.add_paragraph('Category')
doc.add_paragraph(' ├── names (reverse FK) -- members in this category', style='List Bullet')
doc.add_paragraph(' ├── TrainingSpace (reverse FK) -- spaces for this category', style='List Bullet')
doc.add_paragraph(' └── UserProfile (reverse FK) -- trainers in this category', style='List Bullet')

doc.add_page_break()

# Section 13: Sample Data
doc.add_heading('13. Sample Data', 1)

doc.add_heading('13.1 Users (10 accounts)', 2)
user_data = [
    ['user', 'mmm@gmail.com', 'Admin', 'Yes'],
    ['john', 'jj@gmail.com', 'Trainer', 'No'],
    ['kyle', 'kylen@gmail.com', 'Trainee', 'No'],
    ['karen', 'mj@gmail.com', 'Trainee', 'No'],
    ['peter', 'mkll@gmail.com', 'Trainer', 'No'],
    ['jack', 'jack@gmail.com', 'Trainer', 'No'],
    ['yawza', 'yaw@gmail.com', 'Registrar', 'No']
]
add_table_with_style(doc, user_data, ['Username', 'Email', 'Role', 'Superuser'])

doc.add_heading('13.2 Training Categories (4)', 2)
cat_data = [
    ['4', 'Members', 'General membership category'],
    ['5', 'Aerobics', 'High-energy fitness classes'],
    ['6', 'Yoga', 'Mind-body practice'],
    ['14', 'Calisthenics', 'Bodyweight-based strength training']
]
add_table_with_style(doc, cat_data, ['ID', 'Name', 'Description'])

doc.add_heading('13.3 Training Spaces (6)', 2)
space_data = [
    ['Aerobics Studio', 'Aerobics', 'No'],
    ['Aerobics Hall B', 'Aerobics', 'No'],
    ['Calisthenics Park', 'Calisthenics', 'Yes'],
    ['Calisthenics Indoor Gym', 'Calisthenics', 'No'],
    ['Yoga Sanctuary', 'Yoga', 'No'],
    ['Yoga Meditation Dome', 'Yoga', 'No']
]
add_table_with_style(doc, space_data, ['Name', 'Category', 'Maintenance?'])

doc.add_page_break()

# Section 14: Environment Variables
doc.add_heading('14. Environment Variables', 1)
doc.add_paragraph('These must be set in .env file or as environment variables on server:')

env_data = [
    ['DJANGO_SECRET_KEY', 'Cryptographic signing key', 'my-super-secret-key-123'],
    ['DJANGO_DEBUG', 'Enable debug mode (True/False)', 'False (production)'],
    ['ALLOWED_HOSTS', 'Comma-separated domains', 'localhost,myapp.onrender.com'],
    ['DATABASE_URL', 'PostgreSQL connection string', 'postgres://user:pass@host:5432/db'],
    ['EMAIL_HOST_USER', 'Gmail address for emails', 'gym@gmail.com'],
    ['EMAIL_HOST_PASSWORD', 'Gmail app password', 'xxxx xxxx xxxx xxxx'],
    ['EMAIL_HOST', 'SMTP server', 'smtp.gmail.com'],
    ['EMAIL_PORT', 'SMTP port', '587'],
    ['EMAIL_USE_TLS', 'Use TLS encryption', 'true'],
    ['GEMINI_API_KEY', 'Google Gemini API key', 'AIzaSy...'],
    ['CLOUDINARY_CLOUD_NAME', 'Cloudinary cloud name', 'my-cloud'],
    ['CLOUDINARY_API_KEY', 'Cloudinary API key', '123456789'],
    ['CLOUDINARY_API_SECRET', 'Cloudinary API secret', 'secret123']
]
add_table_with_style(doc, env_data, ['Variable', 'Purpose', 'Example'])

# Save document
doc.save(r"C:\Users\natan\OneDrive\Nati's documents\project 103 (2)\SYSTEM_DOCUMENTATION.docx")
print("Word document created successfully!")
