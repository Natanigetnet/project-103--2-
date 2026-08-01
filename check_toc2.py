import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r"C:\Users\natan\OneDrive\Nati's documents\project 103 (2)\Gym_Management_System_Updated_Ch5_6.docx")

# Show exact content of TOC, List of Figures, List of Tables, List of Abbreviations
print("=== TOC (paras 35-59) ===")
for i in range(35, 60):
    p = doc.paragraphs[i]
    print(f'{i}: [{p.style.name}] |{p.text}|')

print("\n=== List of Figures (paras 60-63) ===")
for i in range(60, 64):
    p = doc.paragraphs[i]
    print(f'{i}: [{p.style.name}] |{p.text}|')

print("\n=== List of Tables (paras 64-67) ===")
for i in range(64, 68):
    p = doc.paragraphs[i]
    print(f'{i}: [{p.style.name}] |{p.text}|')

print("\n=== List of Abbreviations (paras 68-75) ===")
for i in range(68, 76):
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        print(f'{i}: [{p.style.name}] |{p.text}|')
