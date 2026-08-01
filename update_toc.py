import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

INPUT_PATH = r"C:\Users\natan\OneDrive\Nati's documents\project 103 (2)\Gym_Management_System_Updated_Ch5_6.docx"
OUTPUT_PATH = r"C:\Users\natan\OneDrive\Nati's documents\project 103 (2)\Gym_Management_System_Updated_Final.docx"

doc = Document(INPUT_PATH)

def add_toc_entry_after(para, text, style_name='Normal (Web)', font_name='Times New Roman', font_size=Pt(12), bold=False):
    """Insert a new paragraph after the given paragraph."""
    parent_element = para._element.getparent()
    new_p = etree.SubElement(parent_element, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
    para._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, para._parent)
    new_para.style = doc.styles[style_name] if style_name in [s.name for s in doc.styles] else doc.styles['Normal']
    run = new_para.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    new_para.paragraph_format.space_before = Pt(2)
    new_para.paragraph_format.space_after = Pt(2)
    new_para.paragraph_format.line_spacing = 1.5
    return new_para

# ============================================================
# 1. Update TABLE OF CONTENTS - add Chapter 5 and 6 entries
# ============================================================
# Find para 57 (last entry of chapter 4) and add after it
toc_insert_after = None
for i, para in enumerate(doc.paragraphs):
    if i == 57:  # Chapter 4 last TOC entry
        toc_insert_after = para
        break

if toc_insert_after is not None:
    # Add entries in reverse order since each inserts right after the anchor
    ch6_entries = [
        '\u2003\u20036.2.3 Recommendations for Similar Projects ..................... 48',
        '\u20036.2.2 Process Recommendations ......................................... 47',
        '\u20036.2.1 System Enhancement Recommendations ........................ 46',
        '6.2 Recommendation ............................................................... 45',
        '6.1 Conclusion ...................................................................... 43',
        'CHAPTER SIX: CONCLUSIONS AND RECOMMENDATIONS .............. 42',
    ]
    ch5_entries = [
        '\u2003\u20035.6.3 Start-Up Strategy .................................................. 41',
        '\u2003\u20035.6.2 Installation .......................................................... 40',
        '\u2003\u20035.6.1 Training .............................................................. 39',
        '\u20035.6 User Manual Preparation .............................................. 39',
        '\u2003\u20035.5.3 Types of Testing Considered ................................. 36',
        '\u2003\u20035.5.2 Testing Tools and Environment ............................. 35',
        '\u2003\u20035.5.1 Test Cases ............................................................ 34',
        '\u20035.5 Testing ........................................................................ 34',
        '\u20035.4 Sample Code ............................................................... 32',
        '\u2003\u20035.3.1 Language Selection Criteria .................................. 31',
        '\u20035.3 Language Specification and Selection Strategy .............. 31',
        '\u20035.2 Hardware and Software Acquisitions ............................. 30',
        '5.1 Introduction ...................................................................... 30',
        'CHAPTER FIVE: IMPLEMENTATION ............................................ 30',
    ]

    all_entries = ch5_entries + ch6_entries + [
        '',
        'REFERENCES .......................................................................... 49',
        'APPENDIX ............................................................................. 51',
    ]

    current = toc_insert_after
    for entry_text in all_entries:
        if entry_text == '':
            # Add empty paragraph
            new_p = etree.SubElement(current._element.getparent(), '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
            current._element.addnext(new_p)
            from docx.text.paragraph import Paragraph
            new_para = Paragraph(new_p, current._parent)
            new_para.paragraph_format.space_before = Pt(0)
            new_para.paragraph_format.space_after = Pt(0)
            current = new_para
        elif entry_text.startswith('CHAPTER') or entry_text.startswith('REFERENCES') or entry_text.startswith('APPENDIX'):
            current = add_toc_entry_after(current, entry_text, style_name='Heading 2', bold=True)
        else:
            current = add_toc_entry_after(current, entry_text, style_name='Normal (Web)')

# ============================================================
# 2. Update LIST OF TABLES - add new tables from Ch5 and Appendix
# ============================================================
lot_insert_after = None
for i, para in enumerate(doc.paragraphs):
    if i == 65:
        lot_insert_after = para
        break

if lot_insert_after is not None:
    new_table_entries = [
        'Table 5.1 Hardware Requirements ............................................. 30',
        'Table 5.2 Software Requirements .............................................. 31',
        'Table 5.3 Language/Framework Comparison Matrix .................. 31',
        'Table 5.4 Test Cases .................................................................. 34',
        'Table 5.5 Testing Tools ............................................................. 35',
        'Table A.1 Database Schema Summary ....................................... 51',
    ]
    current = lot_insert_after
    for entry_text in new_table_entries:
        current = add_toc_entry_after(current, entry_text, style_name='Normal (Web)')

# ============================================================
# 3. Update LIST OF ABBREVIATIONS - add actual abbreviation content
# ============================================================
# Para 68 is "LIST OF ABBREVIATIONS" heading, para 69 is empty
# We need to insert abbreviations after para 69
abbrev_insert_after = None
for i, para in enumerate(doc.paragraphs):
    if i == 69:
        abbrev_insert_after = para
        break

if abbrev_insert_after is not None:
    abbreviations = [
        ('ACT', 'American College of Technology'),
        ('AI', 'Artificial Intelligence'),
        ('API', 'Application Programming Interface'),
        ('BMI', 'Body Mass Index'),
        ('CI/CD', 'Continuous Integration / Continuous Deployment'),
        ('CSRF', 'Cross-Site Request Forgery'),
        ('CSS', 'Cascading Style Sheets'),
        ('DB', 'Database'),
        ('ETB', 'Ethiopian Birr'),
        ('GMS', 'Gym Management System'),
        ('GPS', 'Global Positioning System'),
        ('HTML', 'HyperText Markup Language'),
        ('HTTP', 'HyperText Transfer Protocol'),
        ('HTTPS', 'HyperText Transfer Protocol Secure'),
        ('IDE', 'Integrated Development Environment'),
        ('JS', 'JavaScript'),
        ('MVT', 'Model-View-Template'),
        ('OOSAD', 'Object-Oriented System Analysis and Design'),
        ('ORM', 'Object-Relational Mapping'),
        ('OS', 'Operating System'),
        ('PDF', 'Portable Document Format'),
        ('QR', 'Quick Response'),
        ('RAM', 'Random Access Memory'),
        ('RWD', 'Responsive Web Design'),
        ('SDK', 'Software Development Kit'),
        ('SQL', 'Structured Query Language'),
        ('SSD', 'Solid State Drive'),
        ('SSL', 'Secure Sockets Layer'),
        ('UI', 'User Interface'),
        ('URL', 'Uniform Resource Locator'),
        ('VM', 'Virtual Machine'),
        ('XSS', 'Cross-Site Scripting'),
    ]

    current = abbrev_insert_after
    for abbr, full in abbreviations:
        new_p = etree.SubElement(current._element.getparent(), '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
        current._element.addnext(new_p)
        from docx.text.paragraph import Paragraph
        new_para = Paragraph(new_p, current._parent)
        r1 = new_para.add_run(f'{abbr}')
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12)
        r1.font.bold = True
        r1.font.color.rgb = RGBColor(0, 0, 0)
        r2 = new_para.add_run(f'  -  {full}')
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)
        r2.font.color.rgb = RGBColor(0, 0, 0)
        new_para.paragraph_format.space_before = Pt(2)
        new_para.paragraph_format.space_after = Pt(2)
        new_para.paragraph_format.line_spacing = 1.5
        current = new_para

doc.save(OUTPUT_PATH)
print(f'Document saved: {OUTPUT_PATH}')
print('TOC, List of Tables, and List of Abbreviations updated.')
