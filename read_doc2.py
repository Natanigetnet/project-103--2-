from docx import Document
import sys

try:
    doc = Document(r"C:\Users\natan\OneDrive\Nati's documents\project 103 (2)\Gym_Management_System_123.docx")
    print(f'Successfully opened document')
    print(f'Paragraphs: {len(doc.paragraphs)}')
    print(f'Tables: {len(doc.tables)}')
    print(f'Sections: {len(doc.sections)}')
    
    print('\n=== FIRST 50 PARAGRAPHS ===')
    for i, p in enumerate(doc.paragraphs[:50]):
        text = p.text.strip()
        if text:
            print(f'{i}: [{p.style.name}] {text[:200]}')
        else:
            print(f'{i}: [empty]')
    
    print('\n=== ALL HEADINGS ===')
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith('Heading'):
            print(f'{i}: [{p.style.name}] {p.text}')
    
    print('\n=== TABLE INFO ===')
    for i, table in enumerate(doc.tables):
        print(f'\nTable {i}: {len(table.rows)} rows x {len(table.columns)} columns')
        if len(table.rows) > 0:
            first_row = [cell.text[:40] for cell in table.rows[0].cells]
            print(f'  Headers: {first_row}')
            if len(table.rows) > 1:
                second_row = [cell.text[:40] for cell in table.rows[1].cells]
                print(f'  Row 1: {second_row}')
    
    print('\n=== IMAGES ===')
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1
            print(f'Image: {rel.target_ref}')
    print(f'\nTotal images found: {image_count}')
    
    print('\n=== DOCUMENT PROPERTIES ===')
    print(f'File: Gym_Management_System_123.docx')
    
except Exception as e:
    print(f'Error: {e}')
    import traceback
    traceback.print_exc()
