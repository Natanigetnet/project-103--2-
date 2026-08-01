from docx import Document
import sys

try:
    doc = Document(r"C:\Users\natan\OneDrive\Nati's documents\project 103 (2)\Gym Management System 123 (3).doc")
    print(f'Successfully opened document')
    print(f'Paragraphs: {len(doc.paragraphs)}')
    print(f'Tables: {len(doc.tables)}')
    print(f'Sections: {len(doc.sections)}')
    
    print('\n=== FIRST 30 PARAGRAPHS ===')
    for i, p in enumerate(doc.paragraphs[:30]):
        text = p.text.strip()
        if text:
            print(f'{i}: [{p.style.name}] {text[:150]}')
        else:
            print(f'{i}: [empty]')
    
    print('\n=== TABLE INFO ===')
    for i, table in enumerate(doc.tables):
        print(f'Table {i}: {len(table.rows)} rows x {len(table.columns)} columns')
        if len(table.rows) > 0:
            first_row = [cell.text[:30] for cell in table.rows[0].cells]
            print(f'  Headers: {first_row}')
    
    print('\n=== IMAGES ===')
    # Check for images in relationships
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1
    print(f'Images found: {image_count}')
    
except Exception as e:
    print(f'Error: {e}')
    import traceback
    traceback.print_exc()
