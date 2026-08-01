import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r"C:\Users\natan\OneDrive\Nati's documents\project 103 (2)\Gym_Management_System_Final_With_All.docx")

print("=== FINAL DOCUMENT VERIFICATION ===\n")

print("1. TABLE OF CONTENTS (Chapters 5-6):")
for i, para in enumerate(doc.paragraphs):
    if 'CHAPTER FIVE' in para.text.upper():
        for j in range(i, min(i+25, len(doc.paragraphs))):
            if doc.paragraphs[j].text.strip():
                print(f'  {doc.paragraphs[j].text[:90]}')
        break

print("\n2. LIST OF FIGURES:")
for i, para in enumerate(doc.paragraphs):
    if 'LIST OF FIGURES' in para.text.upper() and para.style.name == 'Heading 1':
        count = 0
        for j in range(i+1, min(i+20, len(doc.paragraphs))):
            if doc.paragraphs[j].text.strip():
                if 'Figure' in doc.paragraphs[j].text:
                    print(f'  {doc.paragraphs[j].text[:90]}')
                    count += 1
                elif count > 0:
                    break
        print(f'  Total: {count} figures')
        break

print("\n3. LIST OF TABLES (new entries from Ch5 & Appendix):")
for i, para in enumerate(doc.paragraphs):
    if 'Table 5.1' in para.text:
        for j in range(max(0, i-1), min(i+7, len(doc.paragraphs))):
            if doc.paragraphs[j].text.strip():
                print(f'  {doc.paragraphs[j].text[:90]}')
        break

print("\n4. LIST OF ABBREVIATIONS:")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == 'LIST OF ABBREVIATIONS' and para.style.name == 'Heading 1':
        count = 0
        for j in range(i+1, min(i+40, len(doc.paragraphs))):
            text = doc.paragraphs[j].text.strip()
            if text and '  -  ' in text:
                print(f'  {text[:70]}')
                count += 1
        print(f'  Total: {count} abbreviations')
        break

print("\n=== DOCUMENT STATISTICS ===")
print(f'Total paragraphs: {len(doc.paragraphs)}')
print(f'Total tables: {len(doc.tables)}')

# Check images preserved
body = doc.element.body
imgdata = body.findall('.//{urn:schemas-microsoft-com:vml}imagedata')
print(f'Images preserved: {len(imgdata)}')

# Count chapters
chapters = []
for para in doc.paragraphs:
    if 'CHAPTER' in para.text.upper() and any(x in para.text.upper() for x in ['ONE', 'TWO', 'THREE', 'FOUR', 'FIVE', 'SIX']):
        chapters.append(para.text[:60])
print(f'\nChapters found: {len(chapters)}')
for ch in chapters:
    print(f'  - {ch}')
