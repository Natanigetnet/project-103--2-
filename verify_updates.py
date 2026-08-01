import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r"C:\Users\natan\OneDrive\Nati's documents\project 103 (2)\Gym_Management_System_Updated_Final.docx")

print("=== Updated TABLE OF CONTENTS ===")
for i in range(35, 90):
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        if p.text.strip():
            print(f'{i}: [{p.style.name}] {p.text[:100]}')

print("\n=== Updated LIST OF TABLES ===")
for i, para in enumerate(doc.paragraphs):
    if 'LIST OF TABLES' in para.text.upper():
        print(f'Found at para {i}')
        for j in range(i, min(i+15, len(doc.paragraphs))):
            if doc.paragraphs[j].text.strip():
                print(f'{j}: {doc.paragraphs[j].text[:100]}')
        break

print("\n=== Updated LIST OF ABBREVIATIONS ===")
for i, para in enumerate(doc.paragraphs):
    if 'LIST OF ABBREVIATIONS' in para.text.upper():
        print(f'Found at para {i}')
        for j in range(i, min(i+40, len(doc.paragraphs))):
            if doc.paragraphs[j].text.strip():
                print(f'{j}: {doc.paragraphs[j].text[:100]}')
        break
