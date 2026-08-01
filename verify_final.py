import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r"C:\Users\natan\OneDrive\Nati's documents\project 103 (2)\Gym_Management_System_Complete.docx")

print("=== TABLE OF CONTENTS (Chapters 4-6) ===")
for i in range(56, 85):
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        if p.text.strip():
            print(f'{i}: {p.text[:100]}')

print("\n=== LIST OF TABLES (new entries) ===")
for i, para in enumerate(doc.paragraphs):
    if 'LIST OF TABLES' in para.text.upper():
        print(f'Found at para {i}')
        for j in range(i, min(i+20, len(doc.paragraphs))):
            if doc.paragraphs[j].text.strip():
                print(f'{j}: {doc.paragraphs[j].text[:100]}')
        break

print("\n=== LIST OF ABBREVIATIONS (first 15 entries) ===")
for i, para in enumerate(doc.paragraphs):
    if 'LIST OF ABBREVIATIONS' in para.text.upper():
        print(f'Found at para {i}')
        count = 0
        for j in range(i, min(i+40, len(doc.paragraphs))):
            if doc.paragraphs[j].text.strip():
                print(f'{j}: {doc.paragraphs[j].text[:80]}')
                count += 1
                if count >= 15:
                    break
        break

print("\n=== Document Statistics ===")
print(f'Total paragraphs: {len(doc.paragraphs)}')
print(f'Total tables: {len(doc.tables)}')

# Check images preserved
body = doc.element.body
imgdata = body.findall('.//{urn:schemas-microsoft-com:vml}imagedata')
print(f'Images preserved: {len(imgdata)}')
