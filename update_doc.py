from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# Open the original document to preserve formatting
original_doc = Document(r"C:\Users\natan\OneDrive\Nati's documents\project 103 (2)\Gym_Management_System_123.docx")

# Create a new document
new_doc = Document()

# Copy all styles from original
for style in original_doc.styles:
    try:
        if style.type is not None:
            new_style = new_doc.styles.add_style(style.name, style.type)
            if style.font:
                new_style.font.name = style.font.name
                new_style.font.size = style.font.size
                new_style.font.bold = style.font.bold
                new_style.font.italic = style.font.italic
            if style.paragraph_format:
                new_style.paragraph_format.alignment = style.paragraph_format.alignment
                new_style.paragraph_format.space_before = style.paragraph_format.space_before
                new_style.paragraph_format.space_after = style.paragraph_format.space_after
                new_style.paragraph_format.line_spacing = style.paragraph_format.line_spacing
    except:
        pass

# Copy sections (page layout)
for section in original_doc.sections:
    new_section = new_doc.add_section()
    new_section.page_width = section.page_width
    new_section.page_height = section.page_height
    new_section.left_margin = section.left_margin
    new_section.right_margin = section.right_margin
    new_section.top_margin = section.top_margin
    new_section.bottom_margin = section.bottom_margin

# Function to copy images from original to new document
def copy_images(doc):
    """Extract and track all images from the document"""
    images = []
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            images.append({
                'rel_id': rel.rId,
                'target_ref': rel.target_ref,
                'image_part': rel.target_part
            })
    return images

# Copy all images
images = copy_images(original_doc)

# Now let's update the content paragraph by paragraph
# We'll keep the same structure but update the text to reflect current system

paragraph_updates = {
    # Update Executive Summary to reflect current features
    71: """The Gym Management System (GMS) has evolved into a comprehensive web-based platform that goes far beyond basic gym administration. Built with Python and Django, the system now incorporates advanced features like QR code-based attendance tracking, AI-powered chat assistance using Google Gemini, a social feed for community engagement, and Telegram integration for real-time notifications. The platform serves four distinct user roles—Admin, Trainer, Trainee, and Registrar—each with specialized dashboards and capabilities. Key features include digital ID cards with QR codes, automated training plan generation with split routines (Push/Pull/Legs, Upper/Lower, etc.), BMI tracking, trainer rating systems, employee payment management, and comprehensive income reporting. The system integrates Chapa payment gateway for Ethiopian users while supporting multiple payment methods. With 20+ database models, 127+ URL endpoints, and 80+ HTML templates, the GMS provides a complete digital ecosystem for modern fitness center management, from member check-in to financial oversight.""",
    
    # Update 1.1 Introduction
    74: """The Gym Management System (GMS) started as a simple web-based solution for fitness center administration but has grown into something much more comprehensive. What began as a basic member registration and scheduling tool has evolved into a full-featured platform that handles everything from QR code attendance tracking to AI-powered chat support. The system now manages four different user types—admins who run the gym, trainers who work with members, trainees who use the facilities, and registrars who handle front desk operations. Each role has its own dashboard with tools specifically designed for their needs. The backend, built with Python and Django, coordinates over 20 different data models that track everything from member body metrics to training session registrations. The frontend uses Bootstrap for responsive design, ensuring the system works well on phones, tablets, and desktops. What makes this system stand out is how it combines practical gym management needs with modern web technologies—like using Google's Gemini AI to answer member questions, generating QR codes for contactless check-ins, and even creating a social feed where members can share workout motivation.""",
    
    # Update 1.2 Background
    76: """When we started this project, most gyms were still using paper logbooks and spreadsheets to manage their operations. Members would sign in on a clipboard, trainers would keep their own notes on paper, and payments were tracked in physical ledgers. This created a lot of problems—data was scattered everywhere, it was easy to lose information, and there was no way to get a real-time picture of what was happening in the gym. We wanted to build something that would bring everything together in one place. The project took on extra importance when we realized how much the payment landscape was changing in Ethiopia. Cash-based systems were becoming harder to manage and audit, so we integrated the Chapa payment gateway, which lets members pay for their memberships online using their preferred local payment methods. This wasn't just about convenience—it meant the system could automatically track when memberships expire, send reminders, and give gym owners accurate financial reports without manual data entry. As the project developed, we kept adding features based on what we saw gyms actually needed: QR codes for fast check-ins, AI chat to answer common questions, training plans that trainers could create and members could follow, and even a social feed to build community among members.""",
    
    # Update 1.3 Statement of the Problem
    78: """The main problem we're solving is that traditional gym management is scattered and manual. When members sign in on paper, you don't know who's actually in the gym right now—which matters for capacity management and safety. When trainers keep their own records, there's no central place to see a member's progress or update their workout plan. When payments are tracked in ledgers, it's easy to lose track of who's paid and whose membership has expired. Security is another big issue—paper records can be seen by anyone, and there's no encryption or access control. Plus, without automated notifications, members forget about upcoming sessions or don't realize their membership is about to expire, which hurts both their progress and the gym's revenue. We also noticed that gyms were missing opportunities to build community among members and provide personalized guidance, which is what led us to add features like the social feed and AI chat support.""",
    
    # Update 1.4.1 General Objective
    81: """The main goal of this project is to build a complete gym management platform that handles everything a modern fitness center needs. We're not just creating a digital version of paper records—we're building a system that actively helps the gym run better. This means automating attendance tracking with QR codes so check-ins are fast and accurate. It means giving trainers tools to create detailed workout plans that members can follow and track. It means integrating payments so memberships are managed automatically. It means using AI to answer member questions and provide guidance. And it means creating a platform where the gym community can interact, share progress, and stay motivated. The system is designed to be secure, easy to use, and scalable—so it can grow with the gym's needs.""",
    
    # Update specific objectives to reflect current features
    83: """Centralized Data Management: The system uses a relational database with 20+ models to store and connect all gym data—members, trainers, categories, attendance logs, payments, training sessions, body metrics, and more. Everything is linked together so you can see a complete picture of any member's journey.""",
    
    84: """Automated Attendance with QR Codes: Instead of manual sign-in sheets, each member gets a digital ID card with a unique QR code. When they arrive at the gym, a registrar or admin scans the code to check them in instantly. The system tracks check-in and check-out times, calculates how long they stayed, and even advances their training split progression automatically.""",
    
    85: """Secure Role-Based Access: The system has four distinct roles (Admin, Trainer, Trainee, Registrar), each with specific permissions. Admins can manage everything, trainers can create sessions and view their trainees, trainees can book sessions and track their progress, and registrars handle check-ins. Each role sees only what they're authorized to access.""",
    
    86: """AI-Powered Communication: The system includes an AI chat assistant powered by Google Gemini that can answer member questions about gym hours, trainer information, how to use the website, and more. It also sends automated notifications for training session reminders, payment confirmations, and trainer assignments through the in-app messaging system.""",
    
    87: """Personalized Training Plans: Trainers can create structured workout programs for their trainees using different split types (Push/Pull/Legs, Upper/Lower, Full Body, etc.). Each day in the plan can have specific exercises with sets, reps, and weights. The system tracks where each trainee is in their program and advances them automatically as they check in to the gym.""",
    
    # Add new objectives for features not in original
    # We'll insert these by modifying the paragraph list
    
    # Update 1.5.1 Technical Feasibility
    93: """The project is technically solid because it's built on Django, which is designed for building secure web applications quickly. Python has a huge ecosystem of libraries that let us add features like QR code generation, image processing, and AI integration without reinventing the wheel. The relational database (SQLite locally, PostgreSQL in production) handles complex queries efficiently—even with hundreds of members and thousands of attendance records. The system is modular, so we can add new features without breaking existing ones. We've also made sure it works well on different devices by using responsive design with Bootstrap. The AI chat feature uses Google's Gemini API, which is reliable and well-documented. Payment integration through Chapa is handled via their API, which provides secure transaction processing. Overall, the technical stack is modern, well-supported, and scalable.""",
    
    # Update 1.5.2 Operational Feasibility
    95: """From a practical standpoint, the system is designed to be easy to use for people with different levels of technical experience. The interface is clean and intuitive, with clear navigation and helpful tooltips. Gym administrators can manage members, view reports, and configure settings through straightforward forms. Trainers can create workout plans and view their trainees' progress without needing to understand the underlying database structure. Members can book sessions, view their training plans, and check their attendance history through a simple dashboard. The QR code check-in system is particularly user-friendly—members just show their phone screen and it's scanned in seconds. We've also added features like the AI chat to help users find answers without needing to contact staff. The system works on phones, tablets, and computers, so people can access it however is most convenient for them.""",
    
    # Update 1.5.3 Economical Feasibility
    97: """Building the GMS makes good financial sense because it uses open-source technologies, which means no expensive software licenses. Django, Python, PostgreSQL—these are all free to use. The main costs are development time and hosting, which are minimal compared to the efficiency gains. The system saves money in several ways: it reduces the time staff spend on manual administration (no more tallying attendance from paper sheets), it prevents revenue loss by automatically tracking membership expirations, and it improves member retention by providing better service. The Chapa integration means payments are processed digitally, reducing the risk of errors or fraud. The QR code system eliminates the need for physical membership cards or key fobs. And the AI chat reduces the burden on staff to answer repetitive questions. Over time, these savings add up to a strong return on investment.""",
    
    # Update 1.6.1 Scope
    105: """The project covers the complete development of a web-based gym management platform. This includes designing the database schema with 20+ models to handle all gym data, building the backend logic in Python/Django to process everything from attendance tracking to payment processing, and creating 80+ responsive HTML templates for the user interface. Specific modules include: member registration with digital ID cards and QR codes, trainer assignment and management, training session creation and booking, QR-based attendance tracking, training plan generation with multiple split types, BMI and body metric tracking, an AI chat assistant, a social feed for community engagement, Telegram integration for notifications, employee payment management, income reporting, and comprehensive admin dashboards. The system supports four user roles with role-specific features and permissions.""",
    
    # Update 1.6.2 Limitations
    107: """The system is a web-based software solution, so it doesn't include physical hardware like biometric scanners or gym equipment. While the QR code attendance system provides the software logic for tracking check-ins, the actual scanning is done using smartphone cameras or webcams—no specialized hardware is required. The system does need an internet connection to function since it's web-based, though we've optimized it to work well even on slower connections. The AI chat feature depends on Google's Gemini API, so it requires internet access and is subject to API rate limits. The payment integration through Chapa is currently configured for Ethiopian users, though the system architecture could support other payment gateways. The social feed and Telegram features are optional and can be disabled if not needed. Overall, the system is designed to be practical and usable in real gym environments without requiring major infrastructure changes.""",
    
    # Update 1.10.1 Frontend Technologies
    123: """The frontend uses standard web technologies—HTML5 for structure, CSS3 for styling, and JavaScript for interactive elements. We're using Bootstrap 5.3 as the responsive framework, which gives us a consistent look across devices and handles the layout grid, forms, and common UI components. Bootstrap Icons provide the iconography throughout the interface. The design follows a glassmorphism aesthetic with backdrop blur effects and a modern color scheme. All templates are server-rendered by Django, which means the pages load quickly and work well even without JavaScript. For interactive features like QR code scanning, we use JavaScript libraries that access the device camera. The AI chat interface uses JavaScript to send messages to the backend and display responses in real-time. The social feed uses AJAX for smooth loading of new posts without page refreshes. Everything is designed to be mobile-first, since most members will access the system from their phones.""",
    
    # Update 1.10.2 Backend Technologies
    125: """The backend is built with Python 3.x and Django 5.2+. Python was chosen because it's readable, has excellent library support, and is well-suited for the kind of data processing we need. Django provides the framework that ties everything together—it handles URL routing, database queries through its ORM (Object-Relational Mapper), user authentication, form validation, and template rendering. The ORM lets us work with database records as Python objects, which makes the code cleaner and easier to maintain. We're using SQLite for local development and PostgreSQL for production, both of which Django supports seamlessly. For image handling, we use Pillow to process uploaded images and generate QR codes. The qrcode library creates the actual QR code images from UUID strings. Google's Gemini API (via the google-genai library) powers the AI chat feature. We use the requests library to integrate with Telegram's Bot API for sending notifications. Email notifications are sent through Django's built-in email system, configured to use Gmail's SMTP server. All sensitive configuration—API keys, database credentials, email passwords—is stored in environment variables using python-dotenv, keeping secrets out of the code.""",
}

# Copy paragraphs from original, updating specific ones
for i, para in enumerate(original_doc.paragraphs):
    # Create new paragraph with same style
    if i in paragraph_updates:
        # Use updated text
        new_para = new_doc.add_paragraph(paragraph_updates[i])
        new_para.style = para.style
    else:
        # Copy original text
        new_para = new_doc.add_paragraph(para.text)
        new_para.style = para.style
    
    # Copy formatting
    if para.alignment:
        new_para.alignment = para.alignment
    
    # Copy font properties for runs
    for orig_run, new_run in zip(para.runs, new_para.runs):
        if orig_run.font.name:
            new_run.font.name = orig_run.font.name
        if orig_run.font.size:
            new_run.font.size = orig_run.font.size
        if orig_run.font.bold:
            new_run.font.bold = orig_run.font.bold
        if orig_run.font.italic:
            new_run.font.italic = orig_run.font.italic

# Copy tables
for table in original_doc.tables:
    new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            new_table.rows[i].cells[j].text = cell.text
            # Copy cell formatting
            for orig_para, new_para in zip(cell.paragraphs, new_table.rows[i].cells[j].paragraphs):
                if orig_para.alignment:
                    new_para.alignment = orig_para.alignment

# Copy images
for rel in original_doc.part.rels.values():
    if "image" in rel.reltype:
        # Add image to new document
        image_part = rel.target_part
        new_doc.part.rels.get_or_add(rel.reltype, image_part)

# Save the new document
output_path = r"C:\Users\natan\OneDrive\Nati's documents\project 103 (2)\Gym_Management_System_Updated.docx"
new_doc.save(output_path)
print(f"Updated document saved to: {output_path}")
