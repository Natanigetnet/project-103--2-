from docx import Document

doc = Document(r"C:\Users\natan\OneDrive\Nati's documents\project 103 (2)\Gym_Management_System_123.docx")

print("=== FULL DOCUMENT CONTENT ===\n")

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text:
        print(f"[{i}] [{p.style.name}] {text}")
    else:
        print(f"[{i}] [empty]")

print("\n=== TABLE CONTENTS ===\n")
for i, table in enumerate(doc.tables):
    print(f"\n--- TABLE {i} ---")
    for row_idx, row in enumerate(table.rows):
        row_text = [cell.text.strip()[:50] for cell in row.cells]
        print(f"Row {row_idx}: {row_text}")
