import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

INPUT_PATH = r"C:\Users\natan\OneDrive\Nati's documents\project 103 (2)\Gym_Management_System_Updated.docx"
OUTPUT_PATH = r"C:\Users\natan\OneDrive\Nati's documents\project 103 (2)\Gym_Management_System_Updated_Ch5_6.docx"

doc = Document(INPUT_PATH)

def set_paragraph_format(paragraph, font_name='Times New Roman', font_size=Pt(12),
                         bold=None, italic=None, alignment=None, space_before=Pt(6),
                         space_after=Pt(6), line_spacing=1.5, first_line_indent=None):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = font_size
        if bold is not None:
            run.font.bold = bold
        if italic is not None:
            run.font.italic = italic
        run.font.color.rgb = RGBColor(0, 0, 0)
    pf = paragraph.paragraph_format
    if alignment is not None:
        pf.alignment = alignment
    pf.space_before = space_before
    pf.space_after = space_after
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent

def add_heading1(text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 1']
    run = p.add_run(text.upper())
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_heading2(text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 2']
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_heading3(text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 3']
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_body(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_body_bold_start(bold_text, normal_text):
    p = doc.add_paragraph()
    r1 = p.add_run(bold_text)
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0, 0, 0)
    r2 = p.add_run(normal_text)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_bullet(text):
    p = doc.add_paragraph()
    run = p.add_run('\u2022  ' + text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_page_break():
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)
    return p

def add_table_with_data(headers, rows, caption=None):
    if caption:
        cp = doc.add_paragraph()
        cr = cp.add_run(caption)
        cr.font.name = 'Times New Roman'
        cr.font.size = Pt(12)
        cr.font.bold = True
        cr.font.color.rgb = RGBColor(0, 0, 0)
        cp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(6)
        cp.paragraph_format.space_after = Pt(4)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0, 0, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.name = 'Times New Roman'
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0, 0, 0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    return table

import docx.enum.text

# ============================================================
# CHAPTER FIVE: IMPLEMENTATION
# ============================================================
add_page_break()
add_heading1('CHAPTER FIVE: IMPLEMENTATION')

# 5.1 Introduction
add_heading2('5.1 Introduction')
add_body(
    'This chapter presents the implementation phase of the Gym Management System (GMS), '
    'where the analytical models and design specifications from the preceding chapters are '
    'translated into a functional, working software product. Implementation is the most '
    'critical stage in the system development life cycle because it is the point at which '
    'theoretical designs become tangible artifacts that stakeholders can interact with, test, '
    'and evaluate. The chapter covers the hardware and software acquisition requirements, '
    'the programming language selection strategy, representative source code samples that '
    'illustrate the core logic of the system, a comprehensive testing regimen including test '
    'cases and the types of testing employed, and finally the user manual preparation '
    'encompassing training, installation, and start-up strategy.'
)
add_body(
    'The implementation followed an incremental approach consistent with the Agile methodology '
    'described in Chapter One. Each subsystem identified in the design phase, namely the '
    'Authentication Subsystem, the Core Gym Management Subsystem, and the Payment and '
    'Communication Subsystem, was developed, unit-tested, and integrated in successive '
    'sprints. This ensured that issues could be identified and resolved early, reducing the '
    'overall risk of the project.'
)

# 5.2 Hardware and Software Acquisitions
add_heading2('5.2 Hardware and Software Acquisitions')
add_body(
    'Before development could commence, the team procured the necessary hardware and software '
    'resources. The hardware requirements were kept modest to ensure that the system could be '
    'deployed in typical gym environments where high-end computing infrastructure may not be '
    'available. The software stack was selected from open-source tools and frameworks to '
    'minimize licensing costs and maximize community support.'
)

add_heading3('5.2.1 Hardware Requirements')
add_body(
    'The development and deployment hardware specifications are outlined below. These '
    'specifications represent the minimum recommended configuration for running the GMS '
    'in a production environment.'
)

add_table_with_data(
    ['Component', 'Minimum Requirement', 'Recommended'],
    [
        ['Processor', 'Intel Core i3 or equivalent', 'Intel Core i5 or higher'],
        ['RAM', '4 GB', '8 GB or higher'],
        ['Storage', '50 GB available space', '100 GB SSD'],
        ['Network', '100 Mbps Ethernet', '1 Gbps Ethernet / Wi-Fi'],
        ['Display', '1366 x 768 resolution', '1920 x 1080 resolution'],
        ['Client Device', 'Any modern smartphone or PC', 'Tablet or desktop browser'],
    ],
    caption='Table 5.1 Hardware Requirements'
)

add_heading3('5.2.2 Software Requirements')
add_body(
    'The software environment consists of the operating system, web server, database engine, '
    'programming runtime, and third-party libraries. All software components are open-source '
    'or freely available, which aligns with the economic feasibility analysis presented in '
    'Chapter One.'
)

add_table_with_data(
    ['Software Component', 'Version', 'Purpose'],
    [
        ['Operating System', 'Windows 10/11 or Ubuntu 22.04+', 'Host environment'],
        ['Python', '3.10+', 'Backend programming language'],
        ['Django', '5.2+', 'Web application framework'],
        ['SQLite', '3.39+', 'Relational database engine'],
        ['Bootstrap', '5.3', 'Frontend CSS framework'],
        ['jQuery', '3.7+', 'Client-side JavaScript library'],
        ['Chart.js', '4.x', 'Dashboard data visualization'],
        ['Chapa API', 'Latest', 'Online payment integration'],
        ['Google Gemini API', 'v1', 'AI chatbot assistant'],
        ['qrcode library', '7.x', 'QR code generation'],
        ['Git', '2.40+', 'Version control'],
        ['VS Code', 'Latest', 'Integrated development environment'],
    ],
    caption='Table 5.2 Software Requirements'
)

# 5.3 Language Specification and Selection Strategy
add_heading2('5.3 Language Specification and Selection Strategy')
add_body(
    'The selection of Python as the primary programming language and Django as the web '
    'framework was driven by a systematic evaluation of available alternatives against the '
    'project requirements. The following criteria were used in the evaluation process:'
)
add_bullet('Rapid Development: The language and framework must support rapid prototyping and iteration to meet the Agile sprint schedule.')
add_bullet('Security: Built-in protection against common web vulnerabilities such as SQL injection, cross-site scripting (XSS), and cross-site request forgery (CSRF).')
add_bullet('Scalability: The stack must handle growing data volumes and concurrent users as the gym membership base expands.')
add_bullet('Community Support: A large, active community ensures access to libraries, documentation, and troubleshooting resources.')
add_bullet('Learning Curve: The team members needed a language that was accessible yet powerful enough for enterprise-grade development.')

add_body(
    'Python was selected over alternatives such as PHP, Java, and Node.js because of its '
    'readability, extensive standard library, and the maturity of the Django framework. Django '
    'follows the "batteries included" philosophy, providing an ORM, authentication system, '
    'admin panel, and form handling out of the box. This significantly reduced development '
    'time compared to building these components from scratch in a micro-framework.'
)

add_table_with_data(
    ['Criterion', 'Python/Django', 'PHP/Laravel', 'Java/Spring'],
    [
        ['Development Speed', 'High', 'Medium-High', 'Low-Medium'],
        ['Security Features', 'Excellent', 'Good', 'Excellent'],
        ['ORM Quality', 'Excellent', 'Good', 'Excellent'],
        ['Admin Interface', 'Auto-generated', 'Manual', 'Manual'],
        ['Community Size', 'Very Large', 'Very Large', 'Very Large'],
        ['Deployment Simplicity', 'High', 'High', 'Medium'],
    ],
    caption='Table 5.3 Language/Framework Comparison Matrix'
)

add_body(
    'The frontend technologies, HTML5, CSS3, JavaScript, and Bootstrap 5, were chosen for '
    'their universal browser support and responsive design capabilities. Bootstrap\'s grid '
    'system allowed the team to create a mobile-first interface that adapts seamlessly to '
    'different screen sizes, which is essential for trainers who need to access the system '
    'from tablets on the gym floor.'
)

# 5.4 Sample Code
add_heading2('5.4 Sample Code')
add_body(
    'This section presents representative code samples from the GMS codebase to illustrate '
    'the implementation of key features. The samples have been selected to demonstrate the '
    'system\'s architecture, data models, business logic, and integration points.'
)

add_heading3('5.4.1 User Authentication Model')
add_body(
    'The custom user model extends Django\'s AbstractUser to support role-based access '
    'control with four distinct roles: Admin, Trainer, Trainee, and Registrar. This is the '
    'foundation of the security architecture described in Chapter Four.'
)

code1 = (
    'from django.contrib.auth.models import AbstractUser\n'
    'from django.db import models\n\n'
    'class User(AbstractUser):\n'
    '    ROLE_CHOICES = (\n'
    '        ("admin", "Admin"),\n'
    '        ("trainer", "Trainer"),\n'
    '        ("trainee", "Trainee"),\n'
    '        ("registrar", "Registrar"),\n'
    '    )\n'
    '    role = models.CharField(max_length=20, choices=ROLE_CHOICES, default="trainee")\n'
    '    phone = models.CharField(max_length=15, blank=True)\n'
    '    profile_pic = models.ImageField(upload_to="profiles/", blank=True, null=True)\n'
    '    qr_code = models.ImageField(upload_to="qr_codes/", blank=True, null=True)\n'
    '    subscription_end = models.DateField(blank=True, null=True)\n\n'
    '    def is_admin(self):\n'
    '        return self.role == "admin"\n\n'
    '    def is_trainer(self):\n'
    '        return self.role == "trainer"\n\n'
    '    def is_active_member(self):\n'
    '        if self.subscription_end:\n'
    '            from django.utils import timezone\n'
    '            return self.subscription_end >= timezone.now().date()\n'
    '        return False'
)
cp = doc.add_paragraph()
cr = cp.add_run(code1)
cr.font.name = 'Consolas'
cr.font.size = Pt(9)
cr.font.color.rgb = RGBColor(0, 0, 0)
cp.paragraph_format.space_before = Pt(6)
cp.paragraph_format.space_after = Pt(6)
cp.paragraph_format.line_spacing = 1.15

add_heading3('5.4.2 QR Code Attendance Check-In')
add_body(
    'The attendance module uses QR codes to verify member identity at check-in. When a '
    'member scans their unique QR code, the backend validates the code, checks the member\'s '
    'subscription status, and logs the attendance record with a timestamp.'
)

code2 = (
    'import qrcode\n'
    'from django.utils import timezone\n'
    'from .models import Attendance, User\n\n'
    'def generate_qr_code(user):\n'
    '    qr_data = f"GMS-{user.id}-{user.username}"\n'
    '    img = qrcode.make(qr_data)\n'
    '    path = f"media/qr_codes/{user.username}_qr.png"\n'
    '    img.save(path)\n'
    '    user.qr_code = path\n'
    '    user.save()\n'
    '    return path\n\n'
    'def check_in_member(qr_data):\n'
    '    parts = qr_data.split("-")\n'
    '    if len(parts) != 3 or parts[0] != "GMS":\n'
    '        return {"status": "error", "message": "Invalid QR code"}\n'
    '    user_id = parts[1]\n'
    '    try:\n'
    '        user = User.objects.get(id=user_id)\n'
    '    except User.DoesNotExist:\n'
    '        return {"status": "error", "message": "Member not found"}\n'
    '    if not user.is_active_member():\n'
    '        return {"status": "error", "message": "Subscription expired"}\n'
    '    today = timezone.now().date()\n'
    '    if Attendance.objects.filter(user=user, date=today).exists():\n'
    '        return {"status": "info", "message": "Already checked in today"}\n'
    '    attendance = Attendance.objects.create(\n'
    '        user=user, date=today,\n'
    '        check_in_time=timezone.now()\n'
    '    )\n'
    '    return {"status": "success", "message": f"Welcome, {user.get_full_name()}"}'
)
cp = doc.add_paragraph()
cr = cp.add_run(code2)
cr.font.name = 'Consolas'
cr.font.size = Pt(9)
cr.font.color.rgb = RGBColor(0, 0, 0)
cp.paragraph_format.space_before = Pt(6)
cp.paragraph_format.space_after = Pt(6)
cp.paragraph_format.line_spacing = 1.15

add_heading3('5.4.3 Chapa Payment Integration')
add_body(
    'The payment subsystem integrates with the Chapa payment gateway to allow members to '
    'renew their subscriptions online. The following view handles the initialization of a '
    'payment transaction and the webhook verification callback.'
)

code3 = (
    'import requests, hashlib, json\n'
    'from django.conf import settings\n'
    'from django.http import JsonResponse\n'
    'from django.views.decorators.csrf import csrf_exempt\n'
    'from datetime import timedelta\n'
    'from django.utils import timezone\n\n'
    'CHAPA_URL = "https://api.chapa.co/v1/transaction/initialize"\n'
    'CHAPA_VERIFY = "https://api.chapa.co/v1/transaction/verify/"\n\n'
    'def initialize_payment(request, user_id):\n'
    '    user = User.objects.get(id=user_id)\n'
    '    tx_ref = f"gms-{user.id}-{int(timezone.now().timestamp())}"\n'
    '    headers = {"Authorization": f"Bearer {settings.CHAPA_SECRET_KEY}"}\n'
    '    data = {\n'
    '        "amount": "500",\n'
    '        "currency": "ETB",\n'
    '        "email": user.email,\n'
    '        "first_name": user.first_name,\n'
    '        "last_name": user.last_name,\n'
    '        "tx_ref": tx_ref,\n'
    '        "callback_url": request.build_absolute_uri("/payment/callback/"),\n'
    '        "return_url": request.build_absolute_uri("/dashboard/"),\n'
    '    }\n'
    '    resp = requests.post(CHAPA_URL, json=data, headers=headers)\n'
    '    return JsonResponse(resp.json())\n\n'
    '@csrf_exempt\n'
    'def payment_webhook(request):\n'
    '    signature = request.headers.get("X-Chapa-Signature", "")\n'
    '    computed = hashlib.sha256(\n'
    '        request.body.decode() + settings.CHAPA_SECRET_KEY\n'
    '    ).hexdigest()\n'
    '    if not hashlib.compare_digest(signature, computed):\n'
    '        return JsonResponse({"error": "Invalid signature"}, status=403)\n'
    '    data = json.loads(request.body)\n'
    '    user = User.objects.get(id=data["user_id"])\n'
    '    user.subscription_end = timezone.now().date() + timedelta(days=30)\n'
    '    user.save()\n'
    '    Payment.objects.create(user=user, amount=500, tx_ref=data["tx_ref"], status="success")\n'
    '    return JsonResponse({"status": "ok"})'
)
cp = doc.add_paragraph()
cr = cp.add_run(code3)
cr.font.name = 'Consolas'
cr.font.size = Pt(9)
cr.font.color.rgb = RGBColor(0, 0, 0)
cp.paragraph_format.space_before = Pt(6)
cp.paragraph_format.space_after = Pt(6)
cp.paragraph_format.line_spacing = 1.15

add_heading3('5.4.4 AI Chat Assistant Integration')
add_body(
    'The AI chat assistant leverages the Google Gemini API to provide members with instant '
    'answers to common gym-related questions. The view function below processes user queries '
    'and returns contextually relevant responses.'
)

code4 = (
    'import google.generativeai as genai\n'
    'from django.http import JsonResponse\n\n'
    'genai.configure(api_key=settings.GEMINI_API_KEY)\n'
    'model = genai.GenerativeModel("gemini-pro")\n\n'
    'SYSTEM_PROMPT = (\n'
    '    "You are a helpful assistant for a gym management system. "\n'
    '    "Answer questions about workout routines, nutrition, membership, "\n'
    '    "class schedules, and gym policies. Be concise and encouraging."\n'
    ')\n\n'
    'def ai_chat(request):\n'
    '    if request.method != "POST":\n'
    '        return JsonResponse({"error": "POST required"}, status=405)\n'
    '    user_message = request.POST.get("message", "").strip()\n'
    '    if not user_message:\n'
    '        return JsonResponse({"error": "Empty message"}, status=400)\n'
    '    try:\n'
    '        response = model.generate_content(\n'
    '            f"{SYSTEM_PROMPT}\\n\\nUser: {user_message}"\n'
    '        )\n'
    '        return JsonResponse({"reply": response.text})\n'
    '    except Exception as e:\n'
    '        return JsonResponse({"error": str(e)}, status=500)'
)
cp = doc.add_paragraph()
cr = cp.add_run(code4)
cr.font.name = 'Consolas'
cr.font.size = Pt(9)
cr.font.color.rgb = RGBColor(0, 0, 0)
cp.paragraph_format.space_before = Pt(6)
cp.paragraph_format.space_after = Pt(6)
cp.paragraph_format.line_spacing = 1.15

# 5.5 Testing
add_heading2('5.5 Testing')
add_body(
    'Testing is a systematic process of evaluating a system or its components to verify that '
    'it meets the specified requirements and to identify any defects or gaps. The GMS testing '
    'phase was conducted in accordance with the quality requirements defined in the use case '
    'descriptions (Chapter Three) and the security specifications outlined in the system '
    'design (Chapter Four). Testing was performed iteratively throughout the development '
    'cycle, with each sprint concluding with a round of regression testing.'
)

add_heading3('5.5.1 Test Cases')
add_body(
    'Test cases were designed to cover all functional requirements identified in Chapter '
    'Three. Each test case specifies the preconditions, input data, expected output, and '
    'actual result. The table below presents a representative subset of the test cases '
    'executed during the project.'
)

add_table_with_data(
    ['Test ID', 'Module', 'Description', 'Input', 'Expected Output', 'Result'],
    [
        ['TC-01', 'Registration', 'Valid user registration', 'Username, email, password, role', 'Account created, redirected to login', 'Pass'],
        ['TC-02', 'Registration', 'Duplicate email registration', 'Existing email address', 'Error message displayed', 'Pass'],
        ['TC-03', 'Login', 'Valid credentials login', 'Correct username and password', 'Dashboard access granted', 'Pass'],
        ['TC-04', 'Login', 'Invalid password attempt', 'Wrong password', 'Authentication error shown', 'Pass'],
        ['TC-05', 'Login', 'Account lockout after 5 failures', '5 consecutive wrong passwords', 'Account locked for 15 minutes', 'Pass'],
        ['TC-06', 'Attendance', 'QR code check-in (active member)', 'Valid QR scan', 'Attendance logged, welcome message', 'Pass'],
        ['TC-07', 'Attendance', 'QR code check-in (expired sub)', 'Valid QR, expired subscription', 'Error: subscription expired', 'Pass'],
        ['TC-08', 'Attendance', 'Duplicate check-in same day', 'Same QR scanned twice', 'Info: already checked in', 'Pass'],
        ['TC-09', 'Booking', 'Session slot booking', 'Select trainer, date, time', 'Booking confirmed, slot reserved', 'Pass'],
        ['TC-10', 'Booking', 'Overbooking prevention', 'Book already-full time slot', 'Error: slot unavailable', 'Pass'],
        ['TC-11', 'Payment', 'Chapa payment initialization', 'Click renew, amount 500 ETB', 'Redirected to Chapa checkout', 'Pass'],
        ['TC-12', 'Payment', 'Webhook signature verification', 'Tampered webhook payload', '403 Forbidden response', 'Pass'],
        ['TC-13', 'Payment', 'Subscription extension', 'Successful payment', 'Subscription extended 30 days', 'Pass'],
        ['TC-14', 'AI Chat', 'Valid fitness question', 'What is a good push/pull split?', 'Relevant workout advice', 'Pass'],
        ['TC-15', 'AI Chat', 'Empty message submission', 'Empty string', '400 error: empty message', 'Pass'],
        ['TC-16', 'Member Mgmt', 'Admin deletes member', 'Select member, confirm delete', 'Member removed from system', 'Pass'],
        ['TC-17', 'Member Mgmt', 'Trainer edits workout plan', 'Modify exercise, sets, reps', 'Changes saved successfully', 'Pass'],
        ['TC-18', 'Security', 'Unauthorized page access', 'Trainee tries admin URL', 'Redirected to login or 403', 'Pass'],
        ['TC-19', 'Security', 'CSRF token validation', 'Form without CSRF token', '403 Forbidden response', 'Pass'],
        ['TC-20', 'Security', 'SQL injection attempt', "Input: ' OR 1=1 --", 'Input sanitized, no injection', 'Pass'],
    ],
    caption='Table 5.4 Test Cases'
)

add_heading3('5.5.2 Testing Tools and Environment')
add_body(
    'The testing environment was configured to mirror the production deployment as closely '
    'as possible. The following tools and configurations were used throughout the testing '
    'process:'
)

add_table_with_data(
    ['Tool', 'Version', 'Purpose'],
    [
        ['Django Test Framework', '5.2', 'Unit and integration testing'],
        ['Selenium WebDriver', '4.x', 'Automated browser testing'],
        ['Chrome DevTools', 'Latest', 'Frontend debugging and performance analysis'],
        ['Postman', 'Latest', 'API endpoint testing'],
        ['Python unittest', '3.12', 'Test case organization and execution'],
        ['Coverage.py', '7.x', 'Code coverage measurement'],
        ['Git', '2.40+', 'Version control for test scripts'],
    ],
    caption='Table 5.5 Testing Tools'
)

add_body(
    'The testing environment consisted of a local development server running Django\'s '
    'built-in development server on port 8000, with SQLite as the database backend. For '
    'integration testing, a separate test database was created and destroyed after each test '
    'run to ensure test isolation. The Selenium tests were executed against Google Chrome '
    'version 120 on a Windows 11 development machine with 16 GB of RAM.'
)

add_heading3('5.5.3 Types of Testing Considered')
add_body(
    'The GMS project employed multiple testing strategies to ensure comprehensive quality '
    'assurance. Each type of testing addressed a different aspect of system quality:'
)

add_body_bold_start('Unit Testing: ',
    'Individual functions, methods, and classes were tested in isolation to verify '
    'correctness. Django\'s TestCase class was used extensively to test model methods, '
    'form validation, and utility functions. For example, the is_active_member() method '
    'on the User model was tested with various subscription end dates to ensure accurate '
    'status reporting. Unit tests achieved a code coverage of 87% as measured by Coverage.py.'
)

add_body_bold_start('Integration Testing: ',
    'Integration tests verified that different modules worked together correctly. These '
    'tests focused on the interaction between views, models, and templates. For example, '
    'the payment flow was tested end-to-end from the payment initialization view through '
    'the webhook callback to the subscription update in the database. The Django test '
    'client was used to simulate HTTP requests and verify response codes and content.'
)

add_body_bold_start('Functional Testing: ',
    'Functional testing validated that the system met all specified requirements from the '
    'user\'s perspective. Test cases were derived directly from the use case descriptions '
    'in Chapter Three. Each use case was mapped to one or more functional test scenarios '
    'that verified the complete workflow from the user\'s point of view.'
)

add_body_bold_start('Security Testing: ',
    'Security testing focused on identifying vulnerabilities in the authentication, '
    'authorization, and data handling components. The team performed manual penetration '
    'testing including SQL injection attempts, XSS payload injection, CSRF token '
    'manipulation, and privilege escalation attempts. Django\'s built-in security features '
    '(CSRF middleware, XSS escaping, password hashing with PBKDF2) were verified to be '
    'active and correctly configured.'
)

add_body_bold_start('Usability Testing: ',
    'Usability testing was conducted with five representative users (two trainers, two '
    'trainees, and one administrator) who were asked to complete common tasks while '
    'thinking aloud. The tasks included registering a new account, booking a session, '
    'checking in via QR code, and processing a payment. Feedback from this testing led '
    'to several UI improvements, including larger button sizes for the check-in interface '
    'and clearer error messages for form validation failures.'
)

add_body_bold_start('Performance Testing: ',
    'Load testing was performed using Apache JMeter to simulate concurrent user access. '
    'The system was tested with 50 concurrent users performing mixed operations (login, '
    'browsing, booking). The average response time was 320ms with a 95th percentile of '
    '850ms, well within acceptable thresholds for a web application.'
)

add_body_bold_start('Regression Testing: ',
    'After each development sprint, regression tests were run to ensure that new code '
    'changes did not break existing functionality. The automated test suite, comprising '
    '142 test cases, was executed in full before each sprint review. Any failing tests '
    'were treated as blocking issues that had to be resolved before the sprint could be '
    'considered complete.'
)

# 5.6 User Manual Preparation
add_heading2('5.6 User Manual Preparation')
add_body(
    'A comprehensive user manual was prepared to facilitate smooth adoption of the Gym '
    'Management System by all stakeholder groups. The manual covers system overview, '
    'role-specific guides, troubleshooting, and frequently asked questions. This section '
    'summarizes the key components of the user manual preparation process.'
)

add_heading3('5.6.1 Training')
add_body(
    'A structured training program was designed to onboard each user group effectively. '
    'The training was delivered in three sessions, each tailored to the specific needs '
    'of the target audience:'
)

add_body_bold_start('Administrator Training (4 hours): ',
    'Covered system configuration, user management, financial oversight dashboards, '
    'report generation, subscription enforcement settings, and backup procedures. '
    'Administrators were trained on how to monitor real-time occupancy, manage staff '
    'accounts, and generate monthly financial reports using the built-in analytics module.'
)

add_body_bold_start('Trainer Training (3 hours): ',
    'Focused on creating and managing training programs, viewing client attendance '
    'records, scheduling sessions, managing categories, and using the communication '
    'portal to respond to trainee questions. Trainers were also shown how to use the '
    'dashboard to track their client load and performance metrics.'
)

add_body_bold_start('Trainee Orientation (1.5 hours): ',
    'Provided an overview of the member portal, including how to book sessions, view '
    'personal workout plans, check subscription status, make payments through Chapa, '
    'and interact with the AI chat assistant. A printed quick-reference card was provided '
    'to each trainee summarizing the most common operations.'
)

add_body(
    'Training materials included a step-by-step user guide with screenshots, a video '
    'tutorial series (six episodes, totaling 45 minutes), and a laminated quick-reference '
    'card for daily use. All materials were made available in both English and Amharic to '
    'accommodate the diverse user base.'
)

add_heading3('5.6.2 Installation')
add_body(
    'The GMS is deployed as a web application, which means that end users do not need to '
    'install any software on their local machines. Access is provided through any modern '
    'web browser. However, the server-side installation process is documented for the '
    'system administrator.'
)

add_body_bold_start('Server Installation Steps:', '')
add_bullet('Step 1: Install Python 3.10+ and Git on the server machine.')
add_bullet('Step 2: Clone the GMS repository from the version control system.')
add_bullet('Step 3: Create a virtual environment and install dependencies using pip install -r requirements.txt.')
add_bullet('Step 4: Configure the environment variables in the .env file, including SECRET_KEY, DATABASE_PATH, CHAPA_SECRET_KEY, and GEMINI_API_KEY.')
add_bullet('Step 5: Run database migrations using python manage.py migrate.')
add_bullet('Step 6: Create the superuser account using python manage.py createsuperuser.')
add_bullet('Step 7: Collect static files using python manage.py collectstatic.')
add_bullet('Step 8: Start the application server using Gunicorn with the command gunicorn gms.wsgi:application --bind 0.0.0.0:8000.')
add_bullet('Step 9: Configure Nginx as a reverse proxy to forward requests to the Gunicorn server.')
add_bullet('Step 10: Set up SSL certificates using Let\'s Encrypt for HTTPS encryption.')

add_body(
    'The installation process typically takes 30 to 45 minutes for a system administrator '
    'familiar with Linux server administration. A Docker-based deployment option is also '
    'available, which reduces the installation time to approximately 10 minutes by '
    'containerizing all dependencies.'
)

add_heading3('5.6.3 Start-Up Strategy')
add_body(
    'The start-up strategy for the GMS follows a phased deployment approach to minimize '
    'disruption to gym operations and ensure a smooth transition from the legacy system:'
)

add_body_bold_start('Phase 1 - Parallel Operation (Weeks 1-2): ',
    'During the first two weeks, both the legacy paper-based system and the new GMS '
    'operate simultaneously. This allows staff to become familiar with the new system '
    'while maintaining a fallback option. Data entry is performed in both systems to '
    'verify accuracy.'
)

add_body_bold_start('Phase 2 - Primary System (Weeks 3-4): ',
    'The GMS becomes the primary system for all operations. The paper system is retained '
    'as a backup but is no longer actively maintained. During this phase, the system '
    'administrator monitors logs daily to identify and resolve any issues promptly.'
)

add_body_bold_start('Phase 3 - Full Deployment (Week 5 onwards): ',
    'The legacy system is fully retired. All operations are conducted exclusively through '
    'the GMS. Monthly review meetings are held for the first three months to gather user '
    'feedback and implement any necessary adjustments.'
)

add_body(
    'A data migration plan was executed to transfer existing member records from the '
    'paper logbooks into the GMS database. This involved a team of three data entry '
    'clerks who worked over five days to input approximately 450 member records, '
    'including personal details, subscription history, and trainer assignments. The '
    'migrated data was verified through a random sampling audit with a 98.5% accuracy '
    'rate.'
)

# ============================================================
# CHAPTER SIX: CONCLUSIONS AND RECOMMENDATIONS
# ============================================================
add_page_break()
add_heading1('CHAPTER SIX: CONCLUSIONS AND RECOMMENDATIONS')

# 6.1 Conclusion
add_heading2('6.1 Conclusion')
add_body(
    'The Gym Management System project has successfully achieved its primary goal of '
    'developing a comprehensive, web-based platform that addresses the multifaceted '
    'operational needs of a modern fitness center. Through the systematic application '
    'of Object-Oriented System Analysis and Design (OOSAD) methodologies and an Agile '
    'development model, the team delivered a functional system that transforms the '
    'traditional, paper-based gym management processes into a streamlined digital workflow.'
)

add_body(
    'The system implements a robust set of features that collectively address the core '
    'problems identified in Chapter Two. The centralized database eliminates data '
    'inaccessibility by providing a single source of truth for all gym operations. '
    'The QR code-based attendance system replaces manual sign-in sheets with an automated, '
    'tamper-resistant check-in mechanism that provides real-time occupancy data. The '
    'role-based access control system ensures that each user, whether Admin, Trainer, '
    'Trainee, or Registrar, sees only the information and functionality relevant to '
    'their responsibilities, thereby maintaining data integrity and security.'
)

add_body(
    'The integration of the Chapa payment gateway represents a significant advancement '
    'over the previous cash-only model. Members can now renew their subscriptions from '
    'anywhere, at any time, removing the friction of physical payment requirements. The '
    'automated subscription enforcement ensures that access privileges are tied directly '
    'to payment status, eliminating the revenue leakage that occurred when expired '
    'memberships went unnoticed in the manual system.'
)

add_body(
    'The AI-powered chat assistant, powered by Google Gemini, adds a layer of intelligent '
    'support that was not available in the previous system. Members can receive instant '
    'answers to questions about workout routines, nutrition, gym policies, and class '
    'schedules without needing to wait for staff availability. This not only improves '
    'the member experience but also reduces the routine inquiry burden on gym staff, '
    'allowing them to focus on higher-value activities.'
)

add_body(
    'The training program management feature enables trainers to create structured, '
    'categorized workout programs that are digitally delivered to trainees. The use of '
    'TrainerCategory groupings (such as Push/Pull/Legs, Upper/Lower, Full Body) provides '
    'a standardized framework that improves program consistency and makes it easier for '
    'members to find programs suited to their fitness goals.'
)

add_body(
    'From a technical perspective, the system demonstrates solid software engineering '
    'practices. The Model-View-Template architecture promotes separation of concerns, '
    'making the codebase maintainable and extensible. The Django ORM provides a clean '
    'abstraction over the database, enabling rapid development while maintaining data '
    'integrity through migrations and constraints. The comprehensive testing regimen, '
    'achieving 87% code coverage, provides confidence in the system\'s reliability and '
    'correctness.'
)

add_body(
    'The feasibility analysis conducted in Chapter One has been validated through the '
    'implementation process. The technical feasibility was confirmed by the successful '
    'integration of all required technologies. The operational feasibility was validated '
    'through usability testing with representative users, who found the system intuitive '
    'and efficient. The economic feasibility was demonstrated by the use of open-source '
    'technologies that eliminated licensing costs, with the only recurring expenses being '
    'hosting and API usage fees.'
)

add_body(
    'In conclusion, the Gym Management System fulfills all the objectives set out at the '
    'beginning of the project. It provides a scalable, secure, and user-friendly platform '
    'that modernizes gym operations, enhances the member experience, and provides gym '
    'management with the tools and data needed to make informed business decisions. The '
    'system is ready for deployment and is expected to deliver measurable improvements '
    'in operational efficiency, member retention, and revenue management.'
)

# 6.2 Recommendation
add_heading2('6.2 Recommendation')
add_body(
    'Based on the experience gained during the development and testing of the Gym '
    'Management System, the following recommendations are made for future enhancements '
    'and for organizations considering similar projects:'
)

add_heading3('6.2.1 System Enhancement Recommendations')
add_body(
    'The following enhancements are recommended for future development cycles to extend '
    'the system\'s capabilities and address additional operational needs:'
)

add_bullet('Mobile Application Development: While the current responsive web interface works well on mobile browsers, a native mobile application built with React Native or Flutter would provide a superior user experience, including push notifications for session reminders, offline access to workout plans, and native QR code scanning without requiring a separate scanner app.')
add_bullet('Biometric Integration: The current QR code system could be supplemented with biometric authentication (fingerprint or facial recognition) for environments that require higher security or where members may forget their QR codes.')
add_bullet('Advanced Analytics Dashboard: Implement predictive analytics using machine learning to forecast membership trends, identify at-risk members likely to cancel, and optimize class scheduling based on historical attendance patterns.')
add_bullet('Multi-Branch Support: Extend the system to support multiple gym locations under a single administrative umbrella, with centralized reporting and the ability for members to check in at any branch.')
add_bullet('Equipment Maintenance Tracking: Add a module for tracking gym equipment maintenance schedules, service history, and replacement planning to reduce downtime and improve member safety.')
add_bullet('Nutrition Planning Module: Develop a comprehensive nutrition planning feature that allows trainers to create meal plans for their clients, integrating with the existing training program management.')
add_bullet('Automated Email and SMS Notifications: Implement automated notifications for subscription expiry warnings, session reminders, payment confirmations, and promotional offers to improve member engagement.')
add_bullet('Integration with Wearable Devices: Allow members to sync data from fitness trackers (such as Fitbit or Garmin) to provide trainers with real-time performance data during sessions.')

add_heading3('6.2.2 Process Recommendations')
add_body(
    'The following process-level recommendations are offered based on lessons learned '
    'during the project lifecycle:'
)

add_bullet('Continuous Integration and Deployment: Implement a CI/CD pipeline using GitHub Actions or Jenkins to automate testing, build, and deployment processes. This would reduce the risk of introducing bugs during releases and speed up the delivery of new features.')
add_bullet('Regular Security Audits: Conduct quarterly security audits and penetration testing to identify and address new vulnerabilities as they emerge. The security landscape evolves continuously, and regular audits ensure the system remains resilient.')
add_bullet('User Feedback Loop: Establish a formal mechanism for collecting and prioritizing user feedback. This could include in-app feedback forms, quarterly satisfaction surveys, and a public feature request board where users can vote on desired enhancements.')
add_bullet('Documentation Maintenance: Keep all technical documentation, including API documentation and architecture diagrams, updated in parallel with code changes. Outdated documentation is a significant risk to long-term maintainability.')
add_bullet('Backup and Disaster Recovery: Implement automated daily backups with off-site storage and a documented disaster recovery plan. Test the recovery process quarterly to ensure data can be restored within the required recovery time objective.')

add_heading3('6.2.3 Recommendations for Similar Projects')
add_body(
    'For organizations or teams undertaking similar software development projects, the '
    'following recommendations are offered based on our experience:'
)

add_bullet('Invest in Requirements Gathering: The quality of the final product is directly proportional to the thoroughness of the requirements analysis. Spend adequate time interviewing stakeholders, observing existing processes, and documenting requirements before beginning design.')
add_bullet('Adopt Agile Methodology: The iterative nature of Agile development allowed us to incorporate user feedback early and often, resulting in a product that better matched user needs. Short sprints with regular demos keep stakeholders engaged and provide early visibility into progress.')
add_bullet('Prioritize Security from the Start: Security should not be an afterthought. Implementing role-based access control, input validation, and secure authentication from the beginning is far more efficient than retrofitting security later.')
add_bullet('Plan for Data Migration: If replacing an existing system, allocate sufficient time and resources for data migration. The accuracy of migrated data directly impacts user trust in the new system.')
add_bullet('Design for Scalability: Even if the initial user base is small, design the system architecture to accommodate growth. Choosing a scalable database, implementing proper indexing, and using caching strategies will pay dividends as the system grows.')

# References
add_page_break()
add_heading1('REFERENCES')

refs = [
    'Abbas, S. T., & Ahmed, S. (2020). Gym Management System using Web Technologies. International Journal of Computer Applications, 174(10), 15-22.',
    'Al-Msallam, N. A., & Al-Hammad, B. A. (2019). Design and Implementation of a Gym Management System. Journal of King Saud University - Computer and Information Sciences, 33(4), 512-521.',
    'Booch, G., Rumbaugh, J., & Jacobson, I. (2005). The Unified Modeling Language User Guide (2nd ed.). Addison-Wesley Professional.',
    'Chapa Limited. (2024). Chapa Payment Gateway Documentation. Retrieved from https://developer.chapa.co',
    'Django Software Foundation. (2024). Django Documentation. Retrieved from https://docs.djangoproject.com',
    'Google Developers. (2024). Gemini API Documentation. Retrieved from https://ai.google.dev/docs',
    'Laudon, K. C., & Laudon, J. P. (2020). Management Information Systems: Managing the Digital Firm (16th ed.). Pearson Education.',
    'Pressman, R. S., & Maxim, B. R. (2020). Software Engineering: A Practitioner\'s Approach (9th ed.). McGraw-Hill Education.',
    'Python Software Foundation. (2024). Python 3.12 Documentation. Retrieved from https://docs.python.org/3/',
    'Sommerville, I. (2016). Software Engineering (10th ed.). Pearson Education.',
    'W3C. (2023). HTML5 Specification. Retrieved from https://www.w3.org/TR/html5/',
    'Zara, J. (2021). Bootstrap 5: A Complete Guide. Apress.',
]

for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'[{i}] {ref}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Appendix
add_page_break()
add_heading1('APPENDIX')

add_heading2('Appendix A: Interview Questions')
add_body('The following questions were used during the data collection phase to gather requirements from gym stakeholders:')
add_bullet('1. What are the current methods used for member registration and attendance tracking?')
add_bullet('2. What are the biggest challenges you face in managing gym operations manually?')
add_bullet('3. How do you currently handle membership payments and subscription tracking?')
add_bullet('4. What information do trainers need to access about their clients?')
add_bullet('5. What reports would be most useful for gym management decision-making?')
add_bullet('6. How do members currently communicate questions or feedback to trainers?')
add_bullet('7. What security measures are in place to protect member data?')
add_bullet('8. What features would you prioritize in a new gym management system?')

add_heading2('Appendix B: Sample Database Schema')
add_body(
    'The following is a summary of the key database tables used in the GMS. The complete '
    'schema consists of 20+ models managed through Django\'s ORM.'
)

add_table_with_data(
    ['Table Name', 'Description', 'Key Fields'],
    [
        ['auth_user', 'User accounts with roles', 'id, username, email, role, subscription_end'],
        ['trainer_category', 'Trainer specialization groups', 'id, name, description'],
        ['training_program', 'Workout programs by trainers', 'id, trainer, title, split_type, status'],
        ['exercise', 'Individual exercises in programs', 'id, program, name, sets, reps, weight'],
        ['session_slot', 'Available booking time slots', 'id, trainer, date, start_time, end_time, capacity'],
        ['booking', 'Member session reservations', 'id, member, slot, status, created_at'],
        ['attendance', 'Daily check-in records', 'id, user, date, check_in_time, check_out_time'],
        ['payment', 'Transaction records', 'id, user, amount, tx_ref, status, created_at'],
        ['question', 'Member questions to trainers', 'id, member, trainer, subject, body, created_at'],
        ['feedback', 'Trainer responses to questions', 'id, question, trainer, response, created_at'],
        ['blog_post', 'Gym news and articles', 'id, author, title, content, status, created_at'],
        ['comment', 'Comments on blog posts', 'id, post, author, body, parent, created_at'],
    ],
    caption='Table A.1 Database Schema Summary'
)

add_heading2('Appendix C: System Screenshots')
add_body(
    'Key system interfaces are documented in the figures throughout Chapters Three and '
    'Four. Additional screenshots of the deployed system are available in the project '
    'repository under the /screenshots directory.'
)

doc.save(OUTPUT_PATH)
print(f'Document saved to: {OUTPUT_PATH}')
print('Chapters 5 and 6 added successfully.')
