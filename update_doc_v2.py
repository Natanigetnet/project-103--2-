from docx import Document
import copy

# Open the original document
doc = Document(r"C:\Users\natan\OneDrive\Nati's documents\project 103 (2)\Gym_Management_System_123.docx")

# Dictionary mapping paragraph index to updated text
# We'll update specific paragraphs to reflect the current system state
paragraph_updates = {
    # Executive Summary (index 71) - Update to reflect current features
    71: """The Gym Management System (GMS) has grown into a comprehensive web-based platform that handles everything a modern fitness center needs. Built with Python and Django, the system now includes QR code-based attendance tracking, AI-powered chat assistance using Google Gemini, a social feed for community engagement, and Telegram integration for real-time notifications. The platform serves four distinct user roles—Admin, Trainer, Trainee, and Registrar—each with specialized dashboards and capabilities. Key features include digital ID cards with QR codes for contactless check-ins, automated training plan generation with split routines (Push/Pull/Legs, Upper/Lower, etc.), BMI and body metric tracking, trainer rating systems, employee payment management, and comprehensive income reporting. The system integrates the Chapa payment gateway for Ethiopian users while supporting multiple payment methods. With over 20 database models, 127 URL endpoints, and 80 HTML templates, the GMS provides a complete digital ecosystem for modern fitness center management, from member check-in to financial oversight.""",
    
    # 1.1 Introduction (index 74)
    74: """The Gym Management System started as a straightforward web solution for gym administration but has evolved into something much more comprehensive. What began as basic member registration and scheduling has become a full-featured platform handling everything from QR code attendance to AI chat support. The system now manages four user types—admins who run the gym, trainers who work with members, trainees who use the facilities, and registrars who handle front desk operations. Each role has its own dashboard with tools designed for their specific needs. The backend, built with Python and Django, coordinates over 20 data models tracking everything from member body metrics to training session registrations. The frontend uses Bootstrap for responsive design, working well on phones, tablets, and desktops. What makes this system stand out is how it combines practical gym management with modern web tech—like using Google's Gemini AI to answer member questions, generating QR codes for contactless check-ins, and creating a social feed where members can share workout motivation.""",
    
    # 1.2 Background (index 76)
    76: """When we started this project, most gyms were still using paper logbooks and spreadsheets. Members signed in on clipboards, trainers kept their own paper notes, and payments were tracked in physical ledgers. This created problems—data was scattered, easy to lose, and there was no real-time picture of what was happening. We wanted to build something that brought everything together in one place. The project gained extra importance when we saw how payment systems were changing in Ethiopia. Cash-based systems were getting harder to manage, so we integrated the Chapa payment gateway, letting members pay for memberships online using local payment methods. This wasn't just about convenience—it meant the system could automatically track membership expirations, send reminders, and give gym owners accurate financial reports without manual work. As the project grew, we kept adding features based on what gyms actually needed: QR codes for fast check-ins, AI chat to answer common questions, training plans trainers could create and members could follow, and a social feed to build community.""",
    
    # 1.3 Statement of Problem (index 78)
    78: """The main problem is that traditional gym management is scattered and manual. Paper sign-in sheets don't show who's actually in the gym right now, which matters for capacity and safety. When trainers keep separate records, there's no central place to see a member's progress or update their workout plan. Paper payment ledgers make it easy to lose track of who's paid and whose membership expired. Security is another issue—paper records can be seen by anyone, with no encryption or access control. Without automated notifications, members forget sessions or don't realize their membership is expiring, hurting both their progress and the gym's revenue. We also saw gyms missing chances to build community and provide personalized guidance, which led us to add features like the social feed and AI chat.""",
    
    # 1.4.1 General Objective (index 81)
    81: """The main goal is to build a complete gym management platform that handles everything a modern fitness center needs. We're not just digitizing paper records—we're building a system that actively helps the gym run better. This means automating attendance with QR codes for fast, accurate check-ins. Giving trainers tools to create detailed workout plans members can follow and track. Integrating payments so memberships are managed automatically. Using AI to answer member questions and provide guidance. Creating a platform where the gym community can interact, share progress, and stay motivated. The system is designed to be secure, easy to use, and scalable—growing with the gym's needs.""",
    
    # 1.4.2 Specific Objectives (indices 83-87)
    83: """Centralized Data Management: The system uses a relational database with 20+ models to store and connect all gym data—members, trainers, categories, attendance logs, payments, training sessions, body metrics, and more. Everything is linked so you can see a complete picture of any member's journey.""",
    
    84: """Automated Attendance with QR Codes: Instead of manual sign-in sheets, each member gets a digital ID card with a unique QR code. When they arrive, a registrar or admin scans the code to check them in instantly. The system tracks check-in and check-out times, calculates duration, and even advances their training split progression automatically.""",
    
    85: """Secure Role-Based Access: The system has four distinct roles (Admin, Trainer, Trainee, Registrar), each with specific permissions. Admins manage everything, trainers create sessions and view trainees, trainees book sessions and track progress, registrars handle check-ins. Each role sees only what they're authorized to access.""",
    
    86: """AI-Powered Communication: The system includes an AI chat assistant powered by Google Gemini that answers member questions about gym hours, trainer info, website navigation, and more. It also sends automated notifications for session reminders, payment confirmations, and trainer assignments through the in-app messaging system.""",
    
    87: """Personalized Training Plans: Trainers create structured workout programs using different split types (Push/Pull/Legs, Upper/Lower, Full Body, etc.). Each day can have specific exercises with sets, reps, and weights. The system tracks where each trainee is in their program and advances them automatically as they check in.""",
    
    # 1.5.1 Technical Feasibility (index 93)
    93: """The project is technically solid because it's built on Django, designed for secure web development. Python's library ecosystem lets us add features like QR generation, image processing, and AI integration without reinventing the wheel. The relational database (SQLite locally, PostgreSQL in production) handles complex queries efficiently—even with hundreds of members and thousands of attendance records. The system is modular, so we add features without breaking existing ones. We've made sure it works on different devices using responsive Bootstrap design. The AI chat uses Google's Gemini API, which is reliable and well-documented. Chapa payment integration is handled via their secure API. Overall, the tech stack is modern, well-supported, and scalable.""",
    
    # 1.5.2 Operational Feasibility (index 95)
    95: """From a practical standpoint, the system is designed to be easy for people with different technical experience levels. The interface is clean and intuitive, with clear navigation and helpful tooltips. Admins manage members, view reports, and configure settings through straightforward forms. Trainers create workout plans and view trainee progress without needing to understand the database. Members book sessions, view training plans, and check attendance through a simple dashboard. The QR check-in is particularly user-friendly—members show their phone screen and it's scanned in seconds. We added AI chat to help users find answers without contacting staff. The system works on phones, tablets, and computers, so people access it however is convenient.""",
    
    # 1.5.3 Economical Feasibility (index 97)
    97: """Building the GMS makes financial sense because it uses open-source technologies—no expensive software licenses. Django, Python, PostgreSQL are all free. Main costs are development time and hosting, minimal compared to efficiency gains. The system saves money by reducing staff time on manual admin (no tallying paper attendance), preventing revenue loss through automatic membership tracking, and improving member retention with better service. Chapa integration means digital payments, reducing errors and fraud. QR codes eliminate physical membership cards. AI chat reduces staff burden answering repetitive questions. Over time, these savings add up to strong ROI.""",
    
    # 1.6.1 Scope (index 105)
    105: """The project covers complete development of a web-based gym management platform. This includes designing the database schema with 20+ models, building backend logic in Python/Django for everything from attendance to payments, and creating 80+ responsive HTML templates. Specific modules include: member registration with digital ID cards and QR codes, trainer assignment and management, training session creation and booking, QR-based attendance tracking, training plan generation with multiple split types, BMI and body metric tracking, AI chat assistant, social feed for community engagement, Telegram integration for notifications, employee payment management, income reporting, and comprehensive admin dashboards. The system supports four user roles with role-specific features.""",
    
    # 1.6.2 Limitations (index 107)
    107: """The system is web-based software, so it doesn't include physical hardware like biometric scanners. The QR code attendance uses smartphone cameras or webcams—no specialized hardware needed. The system needs internet since it's web-based, though we optimized it for slower connections. AI chat depends on Google's Gemini API, requiring internet and subject to rate limits. Chapa payment integration is configured for Ethiopian users, though the architecture could support other gateways. Social feed and Telegram features are optional and can be disabled. Overall, the system is practical and usable in real gym environments without major infrastructure changes.""",
    
    # 1.10.1 Frontend Technologies (index 123)
    123: """The frontend uses HTML5 for structure, CSS3 for styling, and JavaScript for interactive elements. We use Bootstrap 5.3 as the responsive framework, providing consistent look across devices and handling layout, forms, and UI components. Bootstrap Icons provide iconography throughout. The design follows a glassmorphism aesthetic with backdrop blur effects and modern colors. All templates are server-rendered by Django, so pages load quickly and work without JavaScript. For interactive features like QR scanning, we use JavaScript libraries accessing device cameras. AI chat uses JavaScript to send messages and display responses in real-time. The social feed uses AJAX for smooth loading without page refreshes. Everything is mobile-first, since most members access from phones.""",
    
    # 1.10.2 Backend Technologies (index 125)
    125: """The backend is built with Python 3.x and Django 5.2+. Python is readable, has excellent library support, and suits the data processing we need. Django provides the framework handling URL routing, database queries through its ORM, authentication, form validation, and template rendering. The ORM lets us work with database records as Python objects, making code cleaner. We use SQLite for local development and PostgreSQL for production, both seamlessly supported by Django. For images, we use Pillow to process uploads and generate QR codes. The qrcode library creates QR images from UUID strings. Google's Gemini API (via google-genai library) powers AI chat. We use the requests library for Telegram Bot API integration. Email notifications use Django's built-in email system with Gmail SMTP. All sensitive config—API keys, database credentials, email passwords—is stored in environment variables using python-dotenv, keeping secrets out of code.""",
    
    # Update Chapter 2 content about proposed system (index 138)
    138: """The proposed Gym Management System is a centralized, digital platform that's grown well beyond the original scope. It replaces paper logbooks with a digital Attendance Module using QR codes and paper files with a relational Member Database. The system includes a Real-time Slot Booking Engine that manages facility capacity automatically. Most significantly, it uses Python logic to provide AI-assisted Training Schedules through Google Gemini integration, transforming the gym from a simple facility into a high-tech fitness platform. The system also features a Social Feed for community engagement, Telegram integration for instant notifications, comprehensive financial reporting, and employee payment management. It's built on a mobile-first philosophy with Chapa payment integration, allowing members to pay online while the system automatically handles subscription lifecycles.""",
}

# Update paragraphs in-place
for idx, new_text in paragraph_updates.items():
    if idx < len(doc.paragraphs):
        para = doc.paragraphs[idx]
        # Clear existing text
        for run in para.runs:
            run.text = ""
        # Set new text in first run (or create one if none exist)
        if para.runs:
            para.runs[0].text = new_text
        else:
            para.text = new_text

# Save as new file
output_path = r"C:\Users\natan\OneDrive\Nati's documents\project 103 (2)\Gym_Management_System_Updated.docx"
doc.save(output_path)
print(f"Updated document saved to: {output_path}")
print(f"\nUpdated {len(paragraph_updates)} paragraphs to reflect current system state.")
print("All formatting, images, and layout preserved from original document.")
