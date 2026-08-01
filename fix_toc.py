import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
from lxml import etree

INPUT_PATH = r"C:\Users\natan\OneDrive\Nati's documents\project 103 (2)\Gym_Management_System_Updated_Ch5_6.docx"
OUTPUT_PATH = r"C:\Users\natan\OneDrive\Nati's documents\project 103 (2)\Gym_Management_System_Complete.docx"

doc = Document(INPUT_PATH)

def insert_paragraph_after(para, text, style_name='Normal', font_name='Times New Roman', font_size=Pt(12), bold=False):
    """Insert a new paragraph after the given paragraph and return it."""
    # Create new paragraph using the document's add_paragraph
    new_para = doc.add_paragraph()
    
    # Set the style
    if style_name in [s.name for s in doc.styles]:
        new_para.style = doc.styles[style_name]
    
    # Add the text with formatting
    run = new_para.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    new_para.paragraph_format.space_before = Pt(2)
    new_para.paragraph_format.space_after = Pt(2)
    new_para.paragraph_format.line_spacing = 1.5
    
    # Move the paragraph to the correct position (after the reference paragraph)
    para._element.addnext(new_para._element)
    
    return new_para

# ============================================================
# 1. Update TABLE OF CONTENTS
# ============================================================
# Find Chapter 4 TOC entry (para 57)
toc_anchor = doc.paragraphs[57]

if toc_anchor:
    # Add entries in correct order (not reversed)
    toc_entries = [
        ('', 'Normal', False),  # empty line
        ('CHAPTER FIVE: IMPLEMENTATION ............................................ 30', 'Heading 2', True),
        ('5.1 Introduction ...................................................................... 30', 'Normal (Web)', False),
        ('5.2 Hardware and Software Acquisitions ............................. 30', 'Normal (Web)', False),
        ('5.3 Language Specification and Selection Strategy .............. 31', 'Normal (Web)', False),
        ('\u20035.3.1 Language Selection Criteria .................................. 31', 'Normal (Web)', False),
        ('5.4 Sample Code ............................................................... 32', 'Normal (Web)', False),
        ('5.5 Testing ........................................................................ 34', 'Normal (Web)', False),
        ('\u20035.5.1 Test Cases ............................................................ 34', 'Normal (Web)', False),
        ('\u20035.5.2 Testing Tools and Environment ............................. 35', 'Normal (Web)', False),
        ('\u20035.5.3 Types of Testing Considered ................................. 36', 'Normal (Web)', False),
        ('5.6 User Manual Preparation .............................................. 39', 'Normal (Web)', False),
        ('\u20035.6.1 Training .............................................................. 39', 'Normal (Web)', False),
        ('\u20035.6.2 Installation .......................................................... 40', 'Normal (Web)', False),
        ('\u20035.6.3 Start-Up Strategy .................................................. 41', 'Normal (Web)', False),
        ('', 'Normal', False),  # empty line
        ('CHAPTER SIX: CONCLUSIONS AND RECOMMENDATIONS .............. 42', 'Heading 2', True),
        ('6.1 Conclusion ...................................................................... 43', 'Normal (Web)', False),
        ('6.2 Recommendation ............................................................... 45', 'Normal (Web)', False),
        ('\u20036.2.1 System Enhancement Recommendations ........................ 46', 'Normal (Web)', False),
        ('\u20036.2.2 Process Recommendations ......................................... 47', 'Normal (Web)', False),
        ('\u20036.2.3 Recommendations for Similar Projects ..................... 48', 'Normal (Web)', False),
        ('', 'Normal', False),  # empty line
        ('REFERENCES .......................................................................... 49', 'Heading 2', True),
        ('APPENDIX ............................................................................. 51', 'Heading 2', True),
    ]
    
    current = toc_anchor
    for text, style, bold in toc_entries:
        current = insert_paragraph_after(current, text, style_name=style, bold=bold)

# ============================================================
# 2. Update LIST OF TABLES
# ============================================================
# Find the last table entry in the original list (para 65)
last_table_entry = doc.paragraphs[65]

new_tables = [
    'Table 5.1 Hardware Requirements ............................................. 30',
    'Table 5.2 Software Requirements .............................................. 31',
    'Table 5.3 Language/Framework Comparison Matrix .................. 31',
    'Table 5.4 Test Cases .................................................................. 34',
    'Table 5.5 Testing Tools ............................................................. 35',
    'Table A.1 Database Schema Summary ....................................... 51',
]

current = last_table_entry
for table_text in new_tables:
    current = insert_paragraph_after(current, table_text, style_name='Normal (Web)')

# ============================================================
# 3. Update LIST OF ABBREVIATIONS
# ============================================================
# Find the empty paragraph after the heading (para 69)
loa_anchor = doc.paragraphs[69]

abbreviations = [
    ('ACT', 'American College of Technology'),
    ('AI', 'Artificial Intelligence'),
    ('API', 'Application Programming Interface'),
    ('BMI', 'Body Mass Index'),
    ('CI/CD', 'Continuous Integration / Continuous Deployment'),
    ('CSRF', 'Cross-Site Request Forgery'),
    ('CSS', 'Cascading Style Sheets'),
    ('DB', 'Database'),
    ('ETB', 'Ethiopian Birr'),
    ('GMS', 'Gym Management System'),
    ('GPS', 'Global Positioning System'),
    ('HTML', 'HyperText Markup Language'),
    ('HTTP', 'HyperText Transfer Protocol'),
    ('HTTPS', 'HyperText Transfer Protocol Secure'),
    ('IDE', 'Integrated Development Environment'),
    ('JS', 'JavaScript'),
    ('MVT', 'Model-View-Template'),
    ('OOSAD', 'Object-Oriented System Analysis and Design'),
    ('ORM', 'Object-Relational Mapping'),
    ('OS', 'Operating System'),
    ('PDF', 'Portable Document Format'),
    ('QR', 'Quick Response'),
    ('RAM', 'Random Access Memory'),
    ('RWD', 'Responsive Web Design'),
    ('SDK', 'Software Development Kit'),
    ('SQL', 'Structured Query Language'),
    ('SSD', 'Solid State Drive'),
    ('SSL', 'Secure Sockets Layer'),
    ('UI', 'User Interface'),
    ('URL', 'Uniform Resource Locator'),
    ('VM', 'Virtual Machine'),
    ('XSS', 'Cross-Site Scripting'),
]

current = loa_anchor
for abbr, full in abbreviations:
    new_para = doc.add_paragraph()
    
    r1 = new_para.add_run(f'{abbr}')
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0, 0, 0)
    
    r2 = new_para.add_run(f'  -  {full}')
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0, 0, 0)
    
    new_para.paragraph_format.space_before = Pt(2)
    new_para.paragraph_format.space_after = Pt(2)
    new_para.paragraph_format.line_spacing = 1.5
    
    # Move to correct position
    current._element.addnext(new_para._element)
    current = new_para

doc.save(OUTPUT_PATH)
print(f'Document saved: {OUTPUT_PATH}')
print('TOC, List of Tables, and List of Abbreviations updated successfully.')
